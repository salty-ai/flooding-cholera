# -*- coding: utf-8 -*-
"""Build honest real-data cholera manuscript V5 from HoD original base.
(2) real NCDC data + (3) platform/methods framing. No synthetic numbers."""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()
normal = doc.styles['Normal']; normal.font.name='Times New Roman'; normal.font.size=Pt(11)
CENTER=WD_ALIGN_PARAGRAPH.CENTER; JUST=WD_ALIGN_PARAGRAPH.JUSTIFY

def H(t,l=1):
    p=doc.add_heading(t,level=l)
    for r in p.runs: r.font.color.rgb=RGBColor(0,0,0)
    return p
def P(t,italic=False,bold=False,align=JUST,size=11):
    p=doc.add_paragraph(); r=p.add_run(t); r.italic=italic; r.bold=bold; r.font.size=Pt(size)
    p.alignment=align; return p
def caption(t):
    p=doc.add_paragraph(); r=p.add_run(t); r.italic=True; r.font.size=Pt(9)
    p.alignment=CENTER; return p
# map figure-slot -> real image file
import os
from docx.shared import Inches
_IMG={
 "Figure 1":"/root/paper_figures_gen/diag_architecture.png",
 "Figure 2":"/root/paper_figures_gen/diag_pipeline.png",
 "Figure 3":"/root/paper_figures_app/app_dashboard.png",
 "Figure 4":"/root/paper_figures_app/app_map.png",
 "Figure 5":"/root/paper_figures_gen/fig_crossriver_pilot.png",
 "Figure 6":"/root/paper_figures_app/app_facilities.png",
 "Figure 7":"/root/paper_figures_gen/fig_national_burden.png",
 "Figure 8":"/root/paper_figures_gen/fig_lag_signal.png",
 "Figure 9":"/root/paper_figures_app/app_alerts.png",
 "Figure 10":"/root/paper_figures_app/app_agent_generated.png",
 "Figure 11":"/root/paper_figures_gen/diag_roadmap.png",
}
_used=set()
def figph(t, img=None, width=6.5):
    """Embed a real image centered. img overrides keyword mapping."""
    path=img
    if path is None:
        for k,v in _IMG.items():
            if t.startswith(k): path=v; break
    p=doc.add_paragraph(); p.alignment=CENTER
    if path and os.path.exists(path):
        p.add_run().add_picture(path, width=Inches(width))
    else:
        r=p.add_run("[ "+t+" ]"); r.italic=True; r.font.size=Pt(9); r.font.color.rgb=RGBColor(0x88,0x88,0x88)
    return p
def fig(img, width=6.2):
    p=doc.add_paragraph(); p.alignment=CENTER
    if os.path.exists(img): p.add_run().add_picture(img, width=Inches(width))
    return p
def table(headers,rows):
    t=doc.add_table(rows=1,cols=len(headers)); t.style='Light Grid Accent 1'; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; c.text=''; r=c.paragraphs[0].add_run(h); r.bold=True; r.font.size=Pt(9)
    for row in rows:
        cells=t.add_row().cells
        for i,v in enumerate(row):
            cells[i].text=''; r=cells[i].paragraphs[0].add_run(str(v)); r.font.size=Pt(9)
    return t

# ================= TITLE =================
tp=doc.add_paragraph(); tr=tp.add_run("Development and Pilot Validation of a Scalable Earth Observation-Enabled "
"Environmental Health Intelligence Hub for Cholera Surveillance in Nigeria: "
"National Framework and Cross River Sentinel Pilot")
tr.bold=True; tr.font.size=Pt(15); tp.alignment=CENTER
P("Warekuromor Tubolayefa\u00b9*, Adepoju Matthew Olumide\u00b9, Yakubu Tanimu Umar\u00b9",align=CENTER,bold=True)
P("\u00b9 Department of Mission Planning and Satellite Data Management, National Space Research "
  "and Development Agency (NASRDA), Abuja, Nigeria",align=CENTER,italic=True,size=10)
P("* Corresponding author. Head, Mission Planning, Department of Mission Planning and Satellite "
  "Data Management, NASRDA, Abuja, Nigeria.",align=CENTER,italic=True,size=9)

# ================= ABSTRACT =================
H("Abstract",1)
P("This study presents the development, pilot demonstration, and national scale-up design of an "
"Earth Observation (EO)-enabled environmental health surveillance and intelligence hub for cholera in Nigeria. "
"The platform harmonizes epidemiological records with satellite-derived environmental indicators accessed through "
"Google Earth Engine (GEE)—including NASA Global Precipitation Measurement (GPM-IMERG) rainfall, the Normalized "
"Difference Water Index (NDWI) from Sentinel-2, the Normalized Difference Vegetation Index (NDVI) from Landsat, and "
"Synthetic Aperture Radar (SAR) flood extent from Sentinel-1—to characterize environmental drivers of waterborne "
"disease across space and time and to compute an interpretable, LGA-level composite risk score within a geostatistical "
"multi-criteria framework.")
P("The study reports two clearly separated tiers of evidence. At the national tier, the hub compiled officially "
"reported cholera figures from Nigeria Centre for Disease Control and Prevention (NCDC) situation reports, which "
"document 111,062 suspected cases and 3,604 deaths across 33 states and the FCT in 2021—the largest recent national "
"epidemic—followed by substantially lower burdens in 2022–2025. National epidemiological reporting is published at the "
"state level; sub-national LGA-level and ward-level resolution is a design objective of the platform, realized "
"empirically only where line-list data are available. At the sub-national tier, a sentinel pilot for Cross River State "
"ingested a 2021 cholera line-list of 74 suspected cases meeting the national case definition and 4 deaths (case "
"fatality rate 5.4%) across four Local Government Areas (Yakurr, Biase, Calabar Municipal and Bakassi). Of these 74 "
"cases, specimens were obtained for 4, culture was performed for 3, and 1 was culture-confirmed for Vibrio cholerae; "
"the remainder are clinically and epidemiologically defined. The pilot was used to validate dashboard ingestion, "
"spatial joins, and the automated risk-rendering workflow rather than to establish epidemiological effect sizes, and "
"the national figures are explicitly not disaggregated into a sub-national panel.")
P("Observed flood exposure for the four pilot LGAs was derived by intersecting a 2.65-million-polygon dated historical "
"flood archive (2000–2026) with GRID3 administrative boundaries, yielding 191 distinct flood events over the pilot "
"footprint. The platform's AI Surveillance Copilot was evaluated in a controlled schema-detection benchmark of 72 live "
"model invocations over twelve schema variants of the same observed dataset. Under unconstrained prompting the "
"assistant referenced non-existent columns in 97.2% of trials; requiring a file-inspection tool call before dashboard "
"generation eliminated schema hallucination entirely (100% grounding) and raised correct binding of the case-count "
"field from 33.3% to 91.7%. This is reported as a measured capability with an identified failure mode and mitigation, "
"not as a validated epidemiological instrument.")
P("Cross-correlation between satellite-detected surface-water anomalies and reported cholera activity is provided as an "
"exploratory decision-support signal, consistent with the one-to-two-month environmental lag reported in the literature, "
"and is not presented as a validated forecast or as evidence of causation. The results demonstrate the technical "
"feasibility and operational value of an EO-integrated, multi-agency situational-awareness platform, and set out the "
"data-governance and validation steps required to move from retrospective decision support toward prospective early warning.")
P("Keywords: Earth Observation, Cholera Surveillance, Hydro-Epidemiology, Satellite Remote Sensing, Geostatistical Risk "
"Scoring, Flood Risk Mapping, Decision Support Systems, Agentic AI, Nigeria.",italic=True,size=10)

# ================= INTRODUCTION (HoD verbatim) =================
H("1. Introduction",1)
P("Public health surveillance and response is often complicated by the complex interaction between hydrometeorological "
"variables and epidemiological outcomes. Traditional public health surveillance is mostly fragmented and lacks the spatial "
"granularity, as well as the real-time environmental data required to holistically understand the triggers and impact for "
"both proactive and retrospective interventions (Singh et al., 2024; Atobatele et al., 2019). In Nigeria, this challenge is "
"magnified by precarious water insecurity, rapid urban growth, and fragile infrastructure, which collectively intensify the "
"threat of waterborne pathogens.")
P("Environment and climate-induced shifts, such as erratic rainfall, increased temperatures, and amplified flood and drought "
"cycles are altering pathogen ecology and transmission, presenting new challenges to the epidemiology of waterborne diseases "
"like cholera (Akingbola et al., 2025; Madubueze et al., 2025). Cholera (Vibrio cholerae) remains a critical public health "
"challenge in Nigeria and across sub-Saharan Africa, exhibiting severe seasonal spikes driven by localized flooding, poor "
"drainage systems and compromised water, sanitation, and hygiene (WASH) infrastructure. Cholera is primarily characterized by "
"acute watery diarrhoea and extreme dehydration (Madumelu et al., 2026; Kuna and Gajewski, 2017). The effects of climate "
"change consistently alter the epidemiology of waterborne diseases through erratic temperature and precipitation patterns, as "
"well as the increased frequency and intensity of extreme weather events such as flooding and drought. These impacts are "
"consistently escalating the environmental conditions that favour the growth of Vibrio cholerae, thereby amplifying public "
"health risks and broadening transmission windows in vulnerable regions (Yu, 2025; Christaki et al., 2020). Rising temperature "
"patterns have been shown to enable the geographic expansion of cholera-endemic conditions as well as the growth of the "
"bacteria, enabling the pathogen to colonize new areas. Meanwhile, erratic rainfall events and their corresponding impacts, "
"including flash and intense flooding, promote the contamination of water supplies through sewer overflow and latrine "
"inundation, directly and indirectly driving spikes in cholera transmission.")
P("Surveillance and monitoring systems are essential for the early detection of outbreaks within highly vulnerable areas. Early "
"intervention through the provision of clean water and oral cholera vaccines, coupled with improved health services, results in "
"a significant potential reduction in the impact of cholera outbreaks (Mikaberidze, 2025; Chowdhury et al., 2022; Meckawy et "
"al., 2022).")
P("This paper seeks to address this critical operational gap by presenting the conceptualization and pilot deployment of an "
"innovative and bespoke Earth Observation (EO)-enabled Environmental Health Surveillance and Intelligence Hub. The platform is "
"designed to integrate spatially referenced climate and environmental hazard data with epidemiological and infrastructure data "
"layers to provide a unified platform for multi-agency situational awareness. To this end, it transforms the surveillance "
"landscape, replacing irregular, reactionary monitoring with a robust geostatistical decision-support system. This system not "
"only identifies the environmental conditions that precede outbreaks but also delivers granular vulnerability assessments "
"disaggregated to the lowest administrative level for which reliable data exist, enabling resource allocation and response "
"strategies to be tailored to the most at-risk communities within Nigeria, leaving no one behind.")
P("The remainder of this paper is organized as follows. Section 2 reviews Earth Observation for disease surveillance and the "
"environmental determinants of cholera. Section 3 sets out the methodology, system architecture, data provenance, and — in "
"Section 3.1.1 — an explicit statement of the evidence tiers on which the paper's claims rest. Sections 4, 5 and 6 then "
"report those tiers separately and in order: the Cross River sentinel pilot at LGA resolution, the nationally reported "
"burden at state resolution, and a controlled benchmark of the platform's AI assisted-analytics layer. Section 7 discusses "
"the findings and states the limitations, Section 8 sets out the roadmap, and Section 9 concludes. Readers are asked to "
"note that the pilot and the national figures are distinct tiers of evidence and are never combined.")

# ================= 2. EO FOR DISEASE SURVEILLANCE (HoD verbatim lit review) =================
H("2. The Use of Earth Observation (EO) for Disease Surveillance",1)
H("2.1 The Capabilities and Applications of Earth Observation (EO)",2)
P("Earth Observation (EO) refers to the systematic gathering of environmental data about the Earth's physical, chemical, and "
"biological systems through remote sensing technologies, including satellite, aerial, and unmanned aerial vehicle (UAV) "
"platforms (Zhao et al., 2022; Song and Wu, 2021; Cord et al., 2017). These sensors capture a wide spectrum of biophysical "
"variables such as land and water surface temperature, precipitation, soil moisture, vegetation health, atmospheric "
"composition, water quality, land cover, and land use dynamics, without requiring physical contact with the surface. Data are "
"typically stored in raster format, where each pixel represents a geolocated measurement, enabling spatially continuous "
"monitoring across vast regions.")
P("Over the years, EO has evolved from a specialized geospatial scientific tool into a cross-cutting enabler of decision-making "
"across diverse sectors, contributing significantly to sustainable development. It is notably a critical component of "
"interdisciplinary research, with applications spanning health, environmental monitoring, disease surveillance, biodiversity "
"conservation, and habitat characterization (Peprah et al., 2026; Mobasheri, 2022; Avtar et al., 2020; Anderson et al., 2017). "
"This interdisciplinary approach has seen significant successes in the detection, evaluation, and mapping of factors affecting "
"public health, enabling environmental and public health administrations to work collaboratively, improve public health "
"outcomes, disaster risk reduction and efficient resource management. Rather than operating in isolation, EO is most "
"effectively harnessed through cross-sectoral integration, linking environmental monitoring with human health, food security, "
"climate adaptation, and other interrelated domains. For example, vegetation indices and soil moisture data inform drought "
"early warnings, sea surface temperature irregularities help track environmental stress on coastal ecosystems, and "
"satellite-based flood mapping has enabled rapid humanitarian responses during flood incidents.")
P("Critically, EO also plays a central role in monitoring Water, Sanitation, and Hygiene (WASH) infrastructure risks, especially "
"in low-resource settings where conventional monitoring is sparse or absent. By detecting environmental precursors such as "
"flooding, stagnant water accumulation, and temperature shifts, EO provides scalable, near real-time data streams that can "
"inform public health interventions before outbreaks occur, as well as after (Pezanowski et al., 2024; Mobasheri, 2022; Peprah "
"et al., 2026). Thus, EO functions as an environmental health infrastructure that offers a continuous, transboundary data "
"pipeline connecting ecological change to human well-being. Many scientists and resource managers have already recognized the "
"importance of adopting predictive approaches for epidemics and pandemics before their occurrences (Adegoke et al., 2024; "
"Kamalrathne et al., 2023; Morse et al., 2012), and in parallel, the use of EO and GIS approaches has noticeably increased.")

H("2.2 Key Environmental Indicators for Cholera Risk Mapping",2)
P("Various studies have demonstrated that satellite-derived environmental indices are critical for modelling the ecological "
"niche of Vibrio cholerae and identifying high-risk areas before outbreaks occur (Li et al., 2026; Adesina et al., 2026; "
"Kanagaraj and Vijayan, 2024; Madubueze et al., 2025; Akingbola et al., 2025; Bhunia and Shit, 2021; Christaki et al., 2020; "
"Ganesan et al., 2020; Campbell et al., 2020; Escobar et al., 2015). These indices serve as representative measurements for "
"hydrological, thermal, and land surface conditions that promote pathogen persistence and transmission. The selection of "
"indicators in this study is informed by extensive literature, with variables chosen based on their demonstrated predictive "
"value, operational availability, and relevance to Nigeria's diverse ecological zones, from the flood-prone Niger Delta to the "
"drought-vulnerable Sahel.")
P("Precipitation and Hydrological Extremities. Rainfall volume and intensity are primary drivers of surface runoff, sewer "
"overflow, and contamination of drinking water sources (Li et al., 2026; Adesina et al., 2026; Madubueze et al., 2025; "
"Christaki et al., 2020). Previous studies have established that precipitation anomalies, particularly extreme rainfall events, "
"are strongly associated with cholera outbreaks in both coastal and inland settings. Critically, the influence of precipitation "
"on cholera incidence often operates with a lag of one to two months, as heavy rains first accumulate in the environment before "
"contaminating water sources (Armando et al., 2024; Campbell et al., 2020; Reyburn et al., 2011). This lag effect underscores "
"the importance of continuous, near-real-time precipitation monitoring for early warning systems.",bold=False)
P("Surface Water Extent and Flood Persistence. The presence of standing, stagnant surface water bodies is a critical "
"environmental determinant of cholera risk. Vibrio cholerae proliferates in warm, brackish, and nutrient-rich waters, often in "
"association with aquatic zooplankton and phytoplankton (Madumelu et al., 2026; Kuna and Gajewski, 2017). Satellite-derived "
"indices such as the Normalized Difference Water Index (NDWI) have been widely validated for detecting open water features and "
"monitoring flood extent (da Silva Junior et al., 2025; Bhaga et al., 2023; Albertini et al., 2022; McFeeters, 1996). The index "
"is calculated as:")
P("NDWI = (Green − NIR) / (Green + NIR)",align=CENTER,italic=True)
P("Values greater than zero (NDWI > 0) indicate the presence of open water features. In the context of cholera risk mapping, "
"NDWI serves as a proxy for surface water availability and flood persistence—conditions that facilitate human exposure to "
"contaminated water.")
P("Soil Moisture and Vegetation Health. Soil moisture serves as a proxy for environmental humidity and residual moisture "
"retention, conditions that prolong the survival of pathogens (Chowdhury and Rahman, 2025; Campbell et al., 2020). The "
"Normalized Difference Vegetation Index (NDVI), originally developed by Tucker (1979), serves as an indirect optical proxy, "
"assessing soil moisture availability by tracking variations in vegetation health, canopy cover, and biomass density (Yaghobi "
"et al., 2026; Gopo et al., 2026; Singh et al., 2025). The index is calculated as:")
P("NDVI = (NIR − Red) / (NIR + Red)",align=CENTER,italic=True)
P("Higher NDVI values reflect conditions of high humidity and moisture retention that provide optimal breeding grounds for "
"pathogens, increasing their persistence and possible transmission.")
P("Increased Temperature and Climate Change. Laboratory and field studies have demonstrated that Vibrio cholerae growth rates "
"increase with temperature, with optimal conditions occurring in warm, brackish waters (Schets et al., 2025; Brumfield et al., "
"2025). Epidemiological studies have shown that a one-degree increase in temperature can double the risk of a cholera outbreak "
"during rainy seasons, particularly in areas without safe water or sanitation (Gopo et al., 2026; Olagunju et al., 2025; "
"Campbell et al., 2020).")
P("Lag Effects and Temporal Dynamics. A critical consideration in environmental cholera modelling is the presence of lag "
"effects between environmental anomalies and subsequent outbreaks. Campbell et al. (2020) found that chlorophyll-a "
"concentration was most predictive when lagged by one to two months, while land surface temperature was most predictive in the "
"current month. Similarly, Hashizume et al. (2008) found that rainfall effects on cholera incidence peaked at a lag of one to "
"two months, reflecting the time required for environmental contamination to translate into human infection. These findings "
"inform the design of early warning systems, which must account for temporal dynamics to provide actionable lead time.")
P("Machine Learning and Geostatistical Approaches. Several studies have explored machine learning techniques, particularly "
"Random Forest classifiers, for cholera risk prediction using remotely-sensed environmental variables (Singh et al., 2025; "
"Usmani et al., 2023; Campbell et al., 2020). Campbell et al. (2020) achieved a sensitivity of 0.895 for cholera outbreak "
"prediction in coastal India, and Usmani et al. (2023) employed near-real-time remotely sensed data to achieve 92% spatial "
"correspondence between predicted regions and areas of reported cases. However, machine learning approaches require extensive "
"training datasets and may face challenges related to data imbalance, overfitting, and interpretability. The present study "
"adopts a geostatistical risk-scoring framework that complements machine learning by providing interpretable, threshold-based "
"risk indicators that can be operationalized within existing public health workflows.")

H("2.3 The Persistent Gap in Infectious Disease Surveillance",2)
P("Despite these advances, the operational use of EO in infectious disease surveillance—particularly for climate-sensitive, "
"waterborne pathogens like Vibrio cholerae—remains limited and fragmented, especially in high-burden, resource-limited regions "
"(Wright et al., 2025; Singh et al., 2025; Akanda et al., 2011). While decades of research have established strong associations "
"between environmental drivers and cholera outbreaks, this knowledge has not been systematically translated into public health "
"practice. Singh et al. (2025) synthesized 1,137 peer-reviewed studies on remote sensing applications in infectious disease "
"surveillance and found that predictive approaches are significantly more effective than reactive responses. However, they "
"identified persistent challenges that limit operational use, particularly the difficulty of obtaining field-collected "
"environmental data aligned with microbiological measurements, and inconsistencies in epidemiological reporting. These barriers "
"are especially pronounced in settings such as Nigeria, where surveillance systems are often fragmented, data availability is "
"low, and cross-sectoral coordination remains weak.")
P("Yet, despite this well-documented environmental epidemiology, most public health systems in West Africa continue to rely on "
"reactive, case-based surveillance—responding after outbreaks are detected rather than anticipating them. This reactive posture "
"stems from a critical disconnect between the geospatial and health sectors: while EO data can detect environmental risk weeks "
"in advance, these signals rarely reach decision-makers in time or in a usable format.")

H("2.4 The Need for Scalable Cloud-Based Geospatial Systems",2)
P("The use of traditional Geographic Information Systems (GIS) in the analysis of spatial health data is not novel. However, "
"traditional desktop GIS continues to fall short compared to modern EO workflows due to significant limitations in computational "
"capacity, data storage, and real-time processing ability (Geng et al., 2026; Verma and Kotwal, 2025; Lloyd et al., 2020). The "
"volume of satellite data, often measured in petabytes, exceeds the processing capabilities of most local institutions, "
"particularly in low-resource settings.")
P("Cloud-based geospatial platforms offer a transformative alternative. Platforms such as Google Earth Engine (GEE), Amazon Web "
"Services, and the Microsoft Planetary Computer provide instant access to decades of satellite imagery, pre-processed data "
"streams, and massive computing power directly within a web browser (Khachoo et al., 2026; Yu, 2025; Gorelick et al., 2017). "
"When combined with open-source software such as QGIS, SNAP, and web GIS servers, organizations can process, publish, and present "
"spatial data without relying on expensive proprietary software. However, global platforms also raise concerns around data "
"sovereignty, national security, and long-term sustainability, particularly when sensitive health and environmental data are "
"processed or stored abroad. This study advances national data sovereignty by designing an environmental-health surveillance hub "
"that integrates global satellite data with locally sourced observations and incident data, and that is engineered to function "
"even if access to global data is disrupted, by leveraging domestic remote sensing capacity and ground-based monitoring.")

# ================= 3. METHODOLOGY (rewritten, honest + real provenance) =================
H("3. Methodology and System Architecture",1)
H("3.1 Study Setting and Data Architecture",2)
P("This study develops a cholera surveillance system that supports early warning design, retrospective research, and health data "
"reporting. The intended geographic scope covers all 36 states of the Federal Republic of Nigeria and the Federal Capital "
"Territory (FCT), and, where sub-national data are available, the 774 Local Government Areas (LGAs) that constitute "
"Administrative Level 2. The system utilizes publicly available, anonymized epidemiological data from the situation reports and "
"weekly epidemiological reports of the Nigeria Centre for Disease Control and Prevention (NCDC), together with satellite-derived "
"environmental data from the National Space Research and Development Agency (NASRDA) archives, NASA, and the European Space "
"Agency (ESA). The study did not involve human subjects or personally identifiable information.")
P("An important structural characteristic of the national epidemiological data governs the analysis presented here. NCDC cholera "
"situation reports publish suspected cases, deaths, and case fatality ratios primarily at the state level, with LGA-level detail "
"provided only for a subset of high-burden states in a given reporting period. Consequently, this study presents nationally "
"reported figures at the state level as officially published, and treats full 774-LGA and ward-level resolution as a design "
"objective of the platform that is realized empirically only where sub-national line-list data are supplied by state Ministries "
"of Health or the National Cholera Technical Working Group.")

H("3.1.1 Scope and Evidence Tiers",3)
P("Because this paper reports both national figures and a single-state pilot, the two must not be conflated. The national "
"figures in Section 5 are official state-level NCDC counts covering the whole federation; the pilot in Section 4 is a "
"four-LGA technical validation in one state. No national claim in this paper rests on the pilot, and no pilot finding is "
"generalized to the nation. Table 1 states explicitly what is empirically observed, at what resolution, and what remains a "
"design objective. Readers are directed to this table before interpreting any figure in the paper.")
P("Table 1: Evidence tiers — what this study observes, at what resolution, and what remains design intent.",bold=True,size=10)
table(["Tier","Evidence","Spatial resolution","Geographic coverage","Status in this paper"],
[["Tier 1 — National burden","NCDC situation-report counts, 2021–2025","State (Admin-1)","36 states + FCT","Observed; reported as published"],
 ["Tier 2 — Sentinel pilot","Cross River 2021 cholera line-list (74 suspected cases; 1 culture-confirmed)","LGA (Admin-2)","4 LGAs in 1 state","Observed; used for technical validation only"],
 ["Tier 3 — Environmental exposure","Groundsource dated flood-polygon archive intersected with GRID3 boundaries","Polygon, clipped to LGA","4 pilot LGAs (191 events)","Computed from observed archive"],
 ["Tier 4 — Copilot capability","72 live schema-detection trials over 12 schema variants","Not spatial","n/a","Measured on software behaviour, not epidemiology"],
 ["Tier 5 — Design objective","774-LGA / ward-level panel, prospective forecasting, live NCDC–SORMAS sync","LGA / ward","National","NOT achieved; roadmap (Section 7)"]])
P("A national LGA-month panel covering all 774 LGAs was deliberately not constructed. A candidate national panel available to "
"the project was found on inspection to be internally inconsistent with official reporting — it implied 179,253 suspected "
"cases for 2021 against the 111,062 officially reported by NCDC, and attributed Cross River case loads to LGAs that do not "
"appear in the verified state line-list. It was therefore excluded in full, and no synthetic or simulated epidemiological "
"values appear anywhere in this study.")
figph("Figure 1", img="/root/paper_figures_gen/diag_architecture.png")
caption("Figure 1: System architecture — three-tier Earth Observation-enabled cholera surveillance hub, showing raw data "
"sources (Tier 1), backend microservices (Tier 2), the delivery layer (Tier 3), and actionable outputs.")

H("3.2 Data Sources and Provenance",2)
P("To build a georeferenced public health dataset, the hub ingests and harmonizes several data repositories, summarized with "
"their provenance and evidence status in Table 2.")
P("Table 2: Data sources, provenance, and evidence status.",bold=True,size=10)
table(["Data Layer","Source","Resolution / Coverage","Evidence Status"],
[["National cholera burden","NCDC Cholera Situation Reports (Epi Week reports, 2021–2025)","State level; 30–35 states + FCT per year","Observed (official)"],
 ["Cross River pilot line-list","Cross River SMoH / NCDC cholera line-list, 2021","4 LGAs; 74 suspected case records (1 culture-confirmed)","Observed (line-list)"],
 ["Administrative boundaries","GRID3 Nigeria Admin Level 2","774 LGAs, 36 states + FCT","Observed"],
 ["Precipitation","NASA GPM-IMERG via Google Earth Engine","0.1°, national","Computed on demand"],
 ["Surface water (NDWI)","Sentinel-2 MSI via GEE","10–20 m","Computed on demand"],
 ["Vegetation (NDVI)","Landsat-8/9 via GEE","30 m","Computed on demand"],
 ["Flood extent (SAR)","Sentinel-1 GRD via GEE; NEMA flood datasets","10 m; national event-based","Computed / observed"],
 ["Health facilities","FMOH Nigeria Health Facility Registry","46,146 facility records (source count)","Observed (unvalidated)"],
 ["Historical flood polygons","Groundsource flood archive","2000–2026; 2,646,302 dated polygons","Observed (analysed here)"]])
P("Epidemiological registries: nationally reported cholera figures were compiled from NCDC National Cholera Situation Reports. "
"For 2021, NCDC reported 111,062 suspected cases and 3,604 deaths (CFR 3.2%) across 33 states and the FCT (NCDC, 2022). "
"Sub-national pilot data for Cross River State were drawn from a 2021 cholera line-list containing case records with onset "
"dates, epidemiological weeks, wards, settlements, specimen and laboratory fields, and outcomes.")

H("3.2.1 Diagnostic Status of the Pilot Line-List",3)
P("Because diagnostic certainty materially affects how pilot counts should be read, the laboratory fields of the Cross River "
"line-list were audited record by record and are reported in full in Table 3. Of 74 records meeting the case definition, a "
"specimen was recorded as taken for 4 (5.4%), a culture result was recorded for 3 (4.1%), and 1 record (1.4%) was positive "
"for Vibrio cholerae; one culture was negative for enteric pathogens and one was awaiting result at the time of extraction. "
"A rapid diagnostic test result was recorded for 2 records, of which 1 was positive. The remaining 70 records carry no "
"laboratory result and are clinically and epidemiologically defined suspected cases, consistent with standard practice "
"during a declared outbreak when laboratory capacity is saturated.")
P("Accordingly, this paper refers throughout to 74 suspected cases, not to confirmed or culture-referenced cases. This is not "
"a deficiency of the pilot so much as an accurate depiction of the diagnostic environment the platform is designed to serve, "
"and it directly motivates the laboratory-integration priority in Section 7.")
P("Table 3: Laboratory and diagnostic status of the 74 Cross River pilot case records.",bold=True,size=10)
table(["Diagnostic field","Records","Share of 74","Interpretation"],
[["Meets national case definition","74","100.0%","Basis of all pilot counts in this paper"],
 ["Specimen recorded as taken","4","5.4%","Specimen collection was the binding constraint"],
 ["Culture result recorded","3","4.1%","Of which: 1 positive, 1 negative, 1 awaiting"],
 ["Culture-confirmed V. cholerae","1","1.4%","Only laboratory-confirmed case in the pilot"],
 ["RDT result recorded","2","2.7%","1 positive, 1 negative"],
 ["No laboratory result","70","94.6%","Clinically / epidemiologically defined"]])
P("Record hygiene was also addressed during ingestion. The raw line-list contained 77 rows, of which 1 was flagged by the "
"field epidemiologist as not a case and 2 carried no LGA attribution; these 3 rows were excluded. Administrative names "
"appeared in inconsistent forms (for example 'Yakurr' and 'YAKURR'; 'CAL MUNICIPAL' with leading whitespace) and were "
"normalized to GRID3 canonical names before the spatial join. After normalization the ingested totals reconcile exactly to "
"74 cases and 4 deaths across four LGAs, as reported in Table 4.")
P("Geospatial boundaries: administrative boundary shapefiles for Nigeria's second administrative tier (LGA boundaries) were "
"sourced from GRID3. Satellite and climatic data: remote-sensing indices were ingested programmatically via Google Earth Engine, "
"leveraging NASA GPM-IMERG precipitation, Sentinel-2 and Landsat surface hydrology and vegetation dynamics, and Sentinel-1 SAR "
"flood mapping; NASRDA archive data were also ingested.")

H("3.3 Environmental Covariate Processing (Google Earth Engine)",2)
P("Environmental indices were computed within GEE and aggregated to administrative boundaries. Precipitation depth (mm) was "
"derived from GPM-IMERG as a physical driver of surface runoff and contamination overflow. NDWI was computed from Sentinel-2 "
"Green and Near-Infrared bands, with values greater than zero isolating open water, localized pooling, and riverine flood "
"corridors; the flood-mask implementation applies an operational NDWI threshold of 0.3, which is a configurable parameter "
"subject to local calibration rather than a validated universal constant. NDVI was derived from red and near-infrared "
"reflectance as a proxy for canopy moisture and soil saturation. To reflect the one-to-two-month environmental lag reported in "
"the literature (Campbell et al., 2020; Hashizume et al., 2008), the hub aggregates environmental indices over a rolling window "
"preceding each epidemiological reporting period. Where a satellite product is unavailable for a requested location or date, "
"the platform returns an explicit unavailable status rather than substituting synthetic values.")
figph("Figure 2", img="/root/paper_figures_gen/diag_pipeline.png")
caption("Figure 2: End-to-end data processing pipeline — from raw EO satellite grids and NCDC/line-list epidemiological "
"records through preprocessing, spatial join, and geostatistical risk scoring to alerts and decision support.")

H("3.4 Geostatistical Risk Scoring",2)
P("The hub computes a composite cholera vulnerability score for each administrative unit using a normalized weighted "
"multi-criteria decision analysis (MCDA) framework:")
P("R_i = w1·P_i + w2·NDWI_i + w3·NDVI_i + w4·C_hist_i + w5·HF_i",align=CENTER,italic=True)
P("where P_i is the precipitation anomaly, NDWI_i is flood-water persistence, NDVI_i is vegetation/soil-moisture retention, "
"C_hist_i is recent reported case burden, and HF_i is a health-facility proximity index. In the current implementation the "
"weights are fixed, literature-informed heuristic coefficients rather than statistically calibrated parameters; principal "
"component analysis and outcome-based calibration are identified as validation steps in Section 6, not as implemented features. "
"Because the score incorporates recent reported cases as an input, it is an interpretable decision-support signal for current "
"vulnerability and is explicitly not a prospective forecast: any predictive evaluation would require a temporal hold-out design "
"that excludes contemporaneous case counts from the predictors.")

H("3.5 Data Quality, Completeness, and Latency",2)
P("Epidemiological reporting is subject to latency, under-reporting, and variation in diagnostic capacity across states and "
"LGAs. Where sub-national panels are constructed, LGAs with reporting gaps exceeding three consecutive months are flagged, and "
"only units meeting a completeness threshold are retained for longitudinal analysis. Reported figures reflect suspected cases as "
"published by NCDC and should be interpreted as conservative baselines rather than absolute biological counts. These constraints "
"reinforce the value of automated, direct data pipelines to minimize collection latency.")

H("3.6 Observed Flood-Exposure Covariates for the Pilot LGAs",2)
P("Google Earth Engine retrieval is implemented in the platform and is exercised on demand, but the deployment used for this "
"study did not hold an activated Earth Engine entitlement on its host project. Rather than report unavailable covariates, or "
"substitute modelled values, environmental exposure for the pilot LGAs was computed from an independent observed source: a "
"dated historical flood-polygon archive of 2,646,302 records spanning January 2000 to February 2026.")
P("Flood polygons were reprojected to the Africa Albers Equal-Area projection (ESRI:102022) and intersected with GRID3 "
"Admin-2 boundaries for the four pilot LGAs. Two quantities are reported. Event count is the number of distinct dated flood "
"polygons intersecting an LGA. Flooded-area footprint is the geometric union of all such polygons clipped to the LGA "
"boundary; the union is used rather than the sum because flood polygons overlap in time — the same ground floods repeatedly — "
"and summing event areas therefore double-counts and can exceed the area of the LGA itself. The union footprint is bounded "
"above by the LGA area by construction, which provides a built-in validity check on the computation.")
P("A pre-outbreak exposure window of 2011–2020 is additionally reported. This decade precedes the 2021 pilot year and "
"therefore represents information that would have been available to a risk score ex ante, without leakage from the outbreak "
"being described.")

H("3.7 AI Surveillance Copilot: Architecture and Evaluation Design",2)
P("The platform includes a conversational Surveillance Copilot intended to lower the technical barrier for state health "
"officers who must interpret unfamiliar datasets without a data analyst. It is implemented as a streaming, provider-agnostic "
"agent with a bounded tool surface of three functions: query_db, which executes read-only SQL SELECT statements against the "
"surveillance tables and rejects any non-SELECT statement; analyze_file, which performs descriptive pandas operations "
"(head, describe, corr) on an uploaded CSV or Excel file; and generate_ui_spec, which emits a declarative JSON dashboard "
"specification of typed widgets (KPI, chart, map, table) bound to named columns of the uploaded file. The agent runs a "
"bounded tool-use loop of at most ten turns, and the final turn is executed without tools so that a text answer is always "
"produced. In the deployment evaluated here the agent was served by Google Gemini through Vertex AI.")
P("Two properties of this design are relevant to evaluation. First, because generate_ui_spec binds widgets to column names "
"supplied by the model, an incorrect or invented column name produces a dashboard that renders but displays nothing — a "
"silent failure, which is the most dangerous class of failure in an operational surveillance setting. Second, the agent "
"implementation contains an offline fallback path that emits a fixed template dashboard when no provider credential is "
"present. All results reported in Section 6 were generated with live credentials, and every trial was programmatically "
"asserted not to have entered the fallback path; any trial doing so would have been discarded.")
P("Evaluation followed a schema-perturbation design. Twelve variants of the pilot dataset were constructed in which the "
"underlying observed data are identical — the same 74 cases, 4 deaths and 4 LGAs — and only the presentation of the schema "
"changes. This isolates schema-detection capability from data content. Variants span seven difficulty classes: canonical "
"headers, abbreviated NCDC-style headers, dirty headers with stray whitespace and inconsistent casing, decoy schemas "
"containing additional count-like and score-like columns that compete with the true fields, wide tables carrying "
"irrelevant administrative columns, opaque coded headers such as adm2_name and epi_n that carry no domain tokens, and a "
"reduced schema in which the deaths field is absent altogether.")
P("Each variant was evaluated under two prompting conditions. In the unconstrained condition the agent is asked to detect "
"the appropriate columns itself. In the schema-grounded condition the agent is instructed to call analyze_file first to "
"read the true column names, to use only names appearing verbatim in that output, and to omit any widget whose required "
"field does not exist. Each variant–condition pair was repeated three times, giving 72 live invocations. Outcomes were "
"scored automatically against ground-truth semantic roles: schema grounding (whether every column referenced by the "
"emitted specification exists in the file), and correct binding of the LGA, case-count and death-count fields. "
"Infrastructure failures such as HTTP 429 quota responses were detected, retried with exponential backoff, and excluded "
"from scoring so that transport errors could never be recorded as detection failures.")

# ================= 4. PILOT (Tier 2) =================
H("4. Cross River Sentinel Pilot: LGA-Level Technical Validation",1)
P("This section reports the sub-national tier of evidence. Everything in it derives from a single state and four LGAs, and "
"nothing in it should be read as a national result; national figures are reported separately in Section 5.")

H("4.1 Pilot Line-List and Ingestion",2)
P("The hub's ingestion and risk-rendering capability was validated using the Cross River 2021 cholera line-list. After the "
"record-hygiene steps described in Section 3.2.1, the ingested pilot comprised 74 suspected cases and 4 deaths (case fatality "
"rate 5.4%) distributed across four LGAs, with onset spanning epidemiological weeks 6 to 46. One case was culture-confirmed "
"for Vibrio cholerae. The pilot exercised dashboard ingestion, spatial joins to GRID3 LGA boundaries, and automated risk "
"categorization; it is a technical demonstration of the workflow rather than an epidemiological effect estimate. Table 4 "
"summarizes the ingested pilot line-list.")
P("Table 4: Cross River 2021 sentinel pilot line-list summary (suspected cases; 1 culture-confirmed overall).",bold=True,size=10)
table(["Sentinel LGA","State / Zone","Suspected Cases","Deaths","Case Fatality Rate"],
[["Yakurr","Cross River (South-South)","53","0","0.0%"],
 ["Biase","Cross River (South-South)","10","1","10.0%"],
 ["Calabar Municipal","Cross River (South-South)","6","0","0.0%"],
 ["Bakassi","Cross River (South-South)","5","3","60.0%"],
 ["Total","Cross River","74","4","5.4%"]])
P("Two cautions attach to Table 4. First, the LGA-level case fatality ratios are computed on very small denominators: the "
"60.0% ratio for Bakassi reflects 3 deaths among 5 recorded cases and is not a stable estimate of lethality, but rather an "
"indication that only the most severe presentations in that LGA reached a reporting facility. Second, the pilot case "
"fatality rate of 5.4% is not comparable with the 3.2% national ratio reported in Section 5: the former is computed on 74 "
"line-list records from four LGAs, the latter on 111,062 suspected cases reported nationally. The two are different "
"quantities at different tiers and are never combined in this paper.")

H("4.2 Observed Flood Exposure in the Pilot LGAs",2)
P("Intersecting the dated flood archive with the pilot LGA boundaries yielded 191 distinct flood events over the four-LGA "
"footprint across 2000–2026. Table 5 reports event counts and union flooded-area footprints, and Figure 3 visualizes them "
"alongside reported case burden.")
P("Table 5: Observed flood exposure for the four pilot LGAs (Groundsource dated flood archive intersected with GRID3 boundaries).",bold=True,size=10)
table(["Sentinel LGA","LGA area (km²)","Flood events 2000–2026","Years with flooding","Union flooded area (km²)","Union as % of LGA","Pre-outbreak events 2011–2020","Suspected cases 2021"],
[["Yakurr","659.1","18","8","6.11","0.9%","16","53"],
 ["Biase","1,292.7","35","11","1,277.80","98.8%","22","10"],
 ["Calabar Municipal","137.5","136","17","130.71","95.0%","92","6"],
 ["Bakassi","4.2","2","2","4.20","100.0%","2","5"]])
figph("Figure 3", img="/root/paper_figures_gen/fig_pilot_flood.png", width=6.6)
caption("Figure 3: Observed flood exposure and reported burden in the four Cross River pilot LGAs. (a) Count of distinct dated "
"flood events intersecting each LGA, 2000–2026. (b) Cumulative union flooded area for the pre-outbreak decade 2011–2020, "
"expressed as a percentage of LGA area. (c) Pre-outbreak flood exposure against reported 2021 suspected cases. Panel (c) is "
"illustrative only: with four units it supports no inference, and the relationship shown is inverse rather than positive.")
P("The result in Figure 3(c) is reported as observed and is deliberately not presented as supporting the hypothesis. Yakurr, "
"the LGA with by far the largest reported case load (53 of 74), has the lowest flood exposure of the four (0.9% of its area "
"inundated over the pre-outbreak decade), while Bakassi and Biase, which are almost entirely flood-exposed, reported 5 and "
"10 cases respectively. With four spatial units this is not evidence against the environmental hypothesis — it is a sample "
"far too small to test it, and it is strongly confounded by population distribution, health-seeking behaviour, reporting "
"intensity, and the location of the reporting facilities themselves. Calabar Municipal illustrates the confounding directly: "
"it records 136 flood events, by far the most of the four, but this partly reflects that a dense urban LGA is more "
"comprehensively observed than a sparsely populated one.")
P("The methodological point stands independently of the direction of the association. The platform ingested an observed "
"epidemiological line-list and an observed 2.65-million-polygon environmental archive, reconciled them to a common "
"administrative geometry, and produced per-LGA exposure measures that satisfy a physical validity constraint. That is the "
"capability being validated here; establishing an environmental association requires the calibrated, adequately powered "
"design set out in Section 8.")

H("4.3 Dashboard and Map Rendering",2)
P("The dashboard and map successfully rendered geospatial risk layers, LGA-level surveillance reports, and correlation "
"analytics from the ingested line-list, confirming the operational readiness of the ingestion and visualization pipeline. "
"Figures 4 and 5 show the operational interface; Figure 6 shows the pilot case and death distribution derived directly "
"from the ingested line-list.")
figph("Figure 4", img="/root/paper_figures_app/app_dashboard.png")
caption("Figure 4: Main dashboard interface of the surveillance hub, displaying the KPI summary cards, date-range controls, "
"alert level, and the geospatial risk context for the selected surveillance window.")
figph("Figure 5", img="/root/paper_figures_app/app_map.png")
caption("Figure 5: Interactive choropleth risk map (MapLibre/Leaflet) with health-facility overlay, risk-level legend, and "
"time-lapse animation control. The map component is built against the full 774-LGA GRID3 boundary set; in this study it is "
"populated with observed case data only for the four pilot LGAs, and unpopulated units are rendered as no-data rather than "
"as zero risk.")
figph("Figure 6", img="/root/paper_figures_gen/fig_crossriver_pilot.png")
caption("Figure 6: Cross River 2021 sentinel pilot — reported suspected cases and deaths by LGA, derived directly from the "
"ingested line-list (Yakurr, Biase, Calabar Municipal, Bakassi; 74 suspected cases, 4 deaths).")
P("The small case volume and restricted geographic footprint of this pilot preclude inferential conclusions; the exercise "
"establishes technical feasibility and interface usability for public health practitioners. The platform also incorporates "
"the FMOH health-facility registry (Figure 7) to support response planning.")
figph("Figure 7", img="/root/paper_figures_app/app_facilities.png")
caption("Figure 7: FMOH Health Facility Registry overlay — 46,146 source facility records across the 36 states and FCT, with "
"functional-status summary. Record-level validation is a governance step and is not established in this study.")

# ================= 5. NATIONAL (Tier 1) =================
H("5. Nationally Reported Cholera Burden, 2021–2025",1)
P("This section reports the national tier of evidence. All figures are official NCDC state-level counts covering the "
"federation. They are independent of the Cross River pilot in Section 4, are not derived from it, and are not "
"disaggregated by this study into a sub-national panel.")
P("Officially reported cholera figures were compiled from NCDC situation reports (Table 6). The national burden is dominated "
"by the 2021 epidemic — the largest in recent years — followed by markedly lower burdens in 2022–2025, with year-to-year "
"fluctuation driven by rainfall, flooding, displacement, and WASH conditions.")
P("Table 6: Nationally reported cholera burden, Nigeria (NCDC situation reports).",bold=True,size=10)
table(["Year","Suspected Cases","Deaths","CFR","States Reporting","Reporting basis","Source"],
[["2021","111,062","3,604","3.2%","33 states + FCT","Full year, Epi Wk 1–52","NCDC SitRep Epi Wk 52, 2021"],
 ["2022","≈23,550","≥583","2.5%","33 states, 270 LGAs","Cumulative to Epi Wk 47","NCDC SitRep Wks 44–47, 2022"],
 ["2023","Not published as an absolute annual total","—","—","Variable","NCDC reported a reduction of approximately 85% in cases and 79% in deaths relative to 2022","NCDC SitRep Wks 48–52, 2023"],
 ["2024","4,809","156","3.2%","35 states","Cumulative to 21 July 2024","NCDC SitRep 2024"],
 ["2025","1,307","34","2.6%","30 states, 98 LGAs","Cumulative to 20 April 2025","UNICEF/NCDC, 2025"]])
P("Table 6 must be read as a series of officially published reporting snapshots rather than as a set of directly comparable "
"annual totals. The 2021 figure is a completed full-year count, whereas the 2024 and 2025 figures are cumulative "
"year-to-date counts truncated in July and April respectively and will therefore understate their full years. For 2023, "
"NCDC published a proportional reduction rather than an absolute annual total; that relative figure is reported here as "
"published and no absolute value has been imputed for it. Cross-year comparison of these values without accounting for "
"differing reporting windows would be misleading.")
figph("Figure 8", img="/root/paper_figures_gen/fig_national_burden.png")
caption("Figure 8: Nationally reported cholera cases and deaths by year, 2021–2025, compiled from NCDC situation reports. The "
"2021 epidemic (111,062 suspected cases; 3,604 deaths) represents the largest recent national outbreak. Bars for 2024 and "
"2025 reflect part-year cumulative reporting and are not directly comparable with the completed 2021 full-year count.")
P("The 2021 outbreak coincided with above-average rainfall across the Guinea Savannah and Sahel zones, consistent with the "
"environmental determinants reviewed in Section 2. Retrospective overlay of NDWI and flood-extent anomalies against reported "
"cholera activity provides an exploratory decision-support signal for the association between surface-water persistence and "
"subsequent case reporting. Consistent with the literature (Campbell et al., 2020; Hashizume et al., 2008), the environmental "
"signal precedes reported case activity by approximately one to two months; this relationship is presented as a "
"decision-support indicator for prioritization and is not a validated forecast or proof of causation.")
figph("Figure 9", img="/root/paper_figures_gen/fig_lag_signal.png")
caption("Figure 9: Exploratory flood–cholera temporal association (illustrative decision-support signal). Dependence structure "
"and multiplicity are not corrected, and no causal or predictive claim is made.")

# ================= 6. COPILOT (Tier 4) =================
H("6. AI Surveillance Copilot: Capability and Measured Limits",1)
P("This section reports the fourth tier of evidence: a measurement of software behaviour. It makes no epidemiological claim. "
"Its purpose is to establish what the assisted-analytics layer can and cannot reliably do, because an unmeasured AI "
"component in a surveillance platform is an operational liability rather than a feature.")

H("6.1 Rule-Based Alerting and Assisted Analytics",2)
P("Beyond the epidemiological demonstrations, the platform provides a configurable rule-based alert engine (Figure 10) "
"supporting case-surge, high-risk-score and recent-flooding triggers, together with PDF and CSV report export. The alert "
"engine's epidemiological performance — alert sensitivity, specificity and timeliness against observed outbreaks — requires "
"prospective evaluation and is not claimed here.")
figph("Figure 10", img="/root/paper_figures_app/app_alerts.png")
caption("Figure 10: Automated early-warning alert engine — severity and status filters, alert rails, and threshold-rule "
"management for case-surge, high-risk-score, and recent-flooding triggers.")
P("Figure 11 shows the Agent Explorer interface with a dashboard generated by the Surveillance Copilot from the uploaded "
"Cross River 2021 pilot dataset. The specification underlying this dashboard was produced by a live Gemini invocation "
"through Vertex AI and bound its widgets to the LGA, case, death and case-fatality fields of the uploaded file.")
figph("Figure 11", img="/root/paper_figures_app/app_agent_generated.png")
caption("Figure 11: Agent Explorer — an interactive dashboard generated by the Surveillance Copilot (Google Gemini via Vertex "
"AI) from the uploaded Cross River 2021 line-list, with schema binding of the LGA, case, death and case-fatality fields "
"performed by the model rather than configured by hand.")

H("6.2 Schema-Detection Benchmark",2)
P("A single successful demonstration such as Figure 11 establishes that the capability exists; it does not establish that "
"it is reliable. The benchmark described in Section 3.7 was therefore run over twelve schema variants of the same observed "
"dataset under two prompting conditions, with three repetitions each. All 72 invocations were live; none entered the "
"offline fallback path. One trial returned an HTTP 429 quota response, was retried after backoff, and completed "
"successfully. Results are reported in Table 7 and Figure 12.")
P("Table 7: Surveillance Copilot schema-detection performance, 72 live trials (percentages of trials passing).",bold=True,size=10)
table(["Metric","Unconstrained prompt (n=36)","Schema-grounded prompt (n=36)"],
[["Tool invoked (generate_ui_spec called)","100.0%","100.0%"],
 ["Emitted a syntactically valid specification","100.0%","100.0%"],
 ["Schema grounding (no invented column names)","2.8%","100.0%"],
 ["LGA field correctly bound","47.2%","88.9%"],
 ["Case-count field correctly bound","33.3%","91.7%"],
 ["Death-count field correctly bound","24.2%","90.9%"],
 ["Total invented column references","160","0"],
 ["Mean latency per invocation","10.3 s","9.9 s"]])
figph("Figure 12", img="/root/paper_figures_gen/fig_copilot_bench.png", width=6.8)
caption("Figure 12: Surveillance Copilot schema-detection benchmark (72 live Vertex AI invocations). (a) Effect of requiring "
"a file-inspection tool call before dashboard generation, across four correctness metrics. (b) Schema grounding rate by "
"schema difficulty class. Requiring the agent to read the file before describing it eliminates schema hallucination in "
"every difficulty class.")
P("The headline finding is a failure mode, and it is severe. Under unconstrained prompting the assistant emitted at least "
"one non-existent column name in 97.2% of trials — 160 invented column references in total across 36 invocations. The "
"pattern of invention is diagnostic: the most frequently invented names were Latitude and Longitude (26 references each), "
"Date (23), and the canonical epidemiological labels Deaths (17), Cases (15) and LGA (15). The model was not reading the "
"uploaded file; it was reproducing the schema that a cholera surveillance dataset conventionally has. On the opaque "
"variant, whose real columns are adm2_name, epi_n and epi_d, the agent confidently emitted a dashboard bound to LGA, "
"Cases and Deaths — a specification that is syntactically valid, renders without error, and displays nothing at all.")
P("This is precisely the silent-failure class identified in Section 3.7, and it is the reason a single successful "
"demonstration is inadequate evidence for an operational claim. A health officer presented with such a dashboard sees an "
"empty chart, not an error message, and has no indication that the underlying binding was fabricated.")
P("The mitigation is a design change rather than a change of model. Instructing the agent to call analyze_file and read "
"the true column names before generating a specification eliminated schema hallucination completely: 100% grounding "
"across all twelve variants and all seven difficulty classes, with zero invented column references in 36 invocations. "
"Correct binding of the case-count field rose from 33.3% to 91.7%, of the death-count field from 24.2% to 90.9%, and of "
"the LGA field from 47.2% to 88.9%. The improvement carried no latency penalty — mean invocation time was marginally "
"lower at 9.9 seconds against 10.3 seconds — because the additional tool call is offset by shorter deliberation.")
P("Two residual limitations are reported. First, grounding is not the same as correctness: in the dirty-header variant, "
"where column names carry stray leading and trailing whitespace and inconsistent casing, the grounded agent referenced "
"only real columns but still bound the wrong column to the LGA role in all three repetitions, giving 0% role accuracy on "
"that variant despite 100% grounding. Whitespace-bearing headers therefore require deterministic normalization at "
"ingestion rather than reliance on the model. Second, the benchmark measures schema binding on a four-row aggregate "
"dataset from a single disease domain; generalization to large multi-sheet field returns, to other diseases, and to other "
"model providers is not established.")
P("The operational conclusion is that the assisted-analytics layer is fit for use only when constrained to ground itself "
"in the data before describing it, and that the deterministic ingestion path documented in Section 3.2.1 — not the "
"conversational agent — must remain the authoritative route for surveillance data entering the system. The agent's value "
"is in lowering the exploration barrier for non-specialist users, not in serving as a trusted parser.")

# ================= 7. DISCUSSION =================
H("7. Discussion",1)
H("7.1 Operationalizing Intelligence and Multi-Agency Awareness",2)
P("The hub is designed to support after-action review by retrospectively mapping cholera outbreaks against historical "
"environmental data, allowing health officials to identify recurrent high-risk zones and the environmental conditions—such as "
"standing water detected via NDWI—that facilitate pathogen proliferation. In the 2021 national epidemic, retrospective overlay "
"of surface-water anomalies against reported case activity is consistent with an environmental lead time of several weeks; if "
"such signals were flagged operationally in real time, they could support the earlier prepositioning of oral rehydration salts, "
"chlorination supplies, and vaccine stockpiles. This remains a hypothesis for prospective evaluation rather than a demonstrated "
"operational outcome.")
P("A core strength of the platform is the shared operational dashboard, which can provide a single source of truth for disease "
"control agencies, national and state health agencies, and disaster-management teams. By presenting epidemiological and "
"environmental data on a single interface, the hub is intended to bridge the institutional gap between health agencies (e.g. "
"NCDC) and environmental agencies (e.g. NASRDA, NEMA), enabling coordination based on the same spatial intelligence. Realizing "
"this shared awareness in production depends on data-sharing agreements and interoperability arrangements that are governance "
"prerequisites rather than technical outputs of this study.")

H("7.2 Comparison with Existing Systems",2)
P("Table 8 situates the hub relative to existing international and regional climate-health surveillance frameworks across "
"resolution, sensor integration, latency, data sovereignty, and decision support. The hub's distinguishing design goals are "
"sub-national (LGA-level) targeting, multi-sensor integration, sovereign hosting, and an assisted analytics layer. The "
"comparison is one of design intent and architecture; the platforms listed are mature operational systems, whereas the hub "
"described here is validated at pilot scale only, and Table 8 should be read with that asymmetry in mind.")
P("Table 8: Comparative framework of climate-health intelligence platforms.",bold=True,size=10)
table(["Platform","Spatial Resolution","Sensor Integration","Latency","Data Sovereignty","Decision Support","Validation status"],
[["UNICEF cholera platform","Global / Admin-1","Precipitation, temperature","Monthly batch","Foreign cloud","Static risk mapping","Operational"],
 ["WHO Health-Earth","Regional / continental","Precipitation, SST","Bi-weekly / monthly","International servers","Predictive modelling","Operational"],
 ["ECDC EWARS","National / regional","Meteorological grids, land cover","Weekly sitreps","European cloud","Rule-based alerts","Operational"],
 ["NASRDA EO-Hub (this study)","Sub-national (LGA) design; state where reported","GPM-IMERG, Sentinel-2 NDWI, Landsat NDVI, Sentinel-1 SAR, dated flood archive","On-demand retrieval","Sovereign Nigerian infrastructure + NASRDA archives","Rule-based alerts + benchmarked assisted analytics","Pilot only (4 LGAs); not operationally validated"]])

H("7.3 Assisted Analytics as a Surveillance Capability",2)
P("The benchmark in Section 6 has implications beyond this platform. Conversational assistants are being added rapidly to "
"public-health tooling, and their failure modes differ qualitatively from those of conventional software. A schema "
"mis-binding does not raise an exception: it produces a plausible, well-formed, empty dashboard. In a surveillance setting "
"where the user is a district health officer under outbreak conditions rather than a data engineer, that failure is "
"unlikely to be noticed and may be acted upon.")
P("The mitigation demonstrated here is architectural and inexpensive: constrain the agent to inspect the data before "
"describing it, and treat any column name it emits that does not appear verbatim in the source as a hard error rather "
"than a rendering nuisance. This shifts the agent from an authority on the data to a navigator of it. The result also "
"argues for a general reporting norm — that AI components in surveillance platforms should be published with a measured "
"failure rate and a stated mitigation, in the same way a diagnostic assay is published with a sensitivity and a "
"specificity, rather than illustrated with a single successful screenshot.")

H("7.4 Policy Implications",2)
P("The platform is intended to enable three policy shifts: (1) prepositioning of oral rehydration salts, vaccines, and other "
"countermeasures in high-risk LGAs ahead of peak transmission; (2) coordinated inter-agency decision-making on a shared "
"geospatial platform; and (3) monitoring of WASH-intervention effectiveness by tracking subsequent changes in environmental "
"risk indices. LGA-level infrastructure indicators provide a baseline for tracking progress toward WASH targets and identifying "
"investment priorities.")

H("7.5 Limitations",2)
P("This study acknowledges the following limitations, stated explicitly so that the evidence base is not overread.")
P("Diagnostic confirmation. The pilot line-list comprises suspected cases meeting the national case definition. Specimens "
"were obtained for 4 of 74 records and only 1 was culture-confirmed for Vibrio cholerae. The pilot therefore describes "
"reported cholera-like illness, not laboratory-confirmed cholera incidence, and the reported case fatality ratios inherit "
"that uncertainty.")
P("Pilot scope and statistical power. The Cross River demonstration covers four LGAs in one state. Sub-national analyses "
"with four spatial units support no inferential conclusion; the inverse exposure–burden relationship observed in Figure "
"3(c) is reported for transparency and is not interpreted as evidence for or against the environmental hypothesis.")
P("National panel not constructed. Because NCDC publishes at state level, a complete 774-LGA longitudinal panel cannot be "
"constructed from official national reports. A candidate national panel was inspected, found irreconcilable with official "
"NCDC totals, and excluded; nothing in this paper is derived from it. The 774-LGA map component is therefore populated "
"with observed data for four LGAs only.")
P("Earth Engine covariates. GEE retrieval is implemented but was not exercisable on the study host, which lacked an "
"activated Earth Engine entitlement. Precipitation, NDWI and NDVI covariates are consequently described as implemented "
"capability rather than reported as computed results for the pilot; the environmental exposure that is reported derives "
"from the independent dated flood archive. Optical indices, when available, are additionally subject to cloud obstruction "
"during peak rainy season, which Sentinel-1 SAR only partially mitigates (Mobasheri, 2022).")
P("Flood-archive constraints. The archive is a derived product whose per-event detection sensitivity is not characterized "
"here, its coverage density varies with observation frequency — denser in urban Calabar Municipal than in rural Bakassi, "
"which confounds event counts with observability — and it recorded no events within the pilot LGAs during 2021 itself, so "
"exposure is characterized over historical windows rather than contemporaneously with the outbreak.")
P("Risk-score validation. The risk-scoring algorithm uses fixed heuristic weights, incorporates contemporaneous case counts "
"as an input, and has not been prospectively validated against an independent real-time outbreak dataset. The "
"flood–cholera association is exploratory and does not correct for temporal autocorrelation, spatial dependence, or "
"multiple comparisons.")
P("Copilot benchmark scope. The benchmark evaluates schema binding on small aggregate tables from one disease domain "
"against a single model provider. Grounding does not guarantee correct role assignment, as the dirty-header variant "
"demonstrates. Generalization to large multi-sheet field returns, other diseases, and other providers is not established.")
P("Integration and sustainability. Live NCDC/SORMAS synchronization and multi-provider AI operation are design features and "
"roadmap items, not evidenced production integrations. Health-facility records are ingested as supplied and are not "
"record-level validated. Operation requires stable connectivity and technical skills that may not be uniformly available "
"in state health offices, and long-term sustainability depends on institutional commitment, funding, and data-governance "
"agreements.")

# ================= 8. FUTURE ROADMAP =================
H("8. Future Roadmap",1)
P("The evolution of the hub encompasses six priorities. First, laboratory-confirmation linkage: integrating culture and "
"rapid-diagnostic results from the national reference laboratory network so that confirmed and suspected cases can be "
"distinguished at ingestion rather than reconstructed by audit, directly addressing the 1-in-74 confirmation rate "
"documented in Section 3.2.1. Second, socio-demographic vulnerability layer ingestion: overlaying national "
"health-infrastructure registries and Demographic and Health Survey (DHS) microdata—including access to safely managed "
"drinking water, improved sanitation, and household wealth quintiles—so that an environmental hazard triggers a high-risk "
"alert only where it intersects weak WASH infrastructure. Third, statistical calibration and validation: replacing fixed "
"heuristic weights with principal-component or outcome-calibrated weights, and evaluating the risk score on a temporal "
"hold-out with discrimination and calibration metrics, dependence-aware inference, and geographic leave-one-region-out "
"sensitivity analysis, across an adequately powered multi-state sentinel network rather than four LGAs. Fourth, "
"programmatic NCDC surveillance synchronization: establishing a governed API linkage (including SORMAS interoperability) "
"to replace manual extraction with automated case and death ingestion. Fifth, sovereign EO expansion: activating Earth "
"Engine entitlement on the production host, and integrating NASRDA historical archives and the planned four-satellite "
"constellation (three optical, one SAR) to reduce dependence on foreign data and improve cloud-penetrating flood mapping. "
"Sixth, hardened assisted analytics: enforcing the schema-grounding constraint validated in Section 6 as a system-level "
"invariant rather than a prompt convention, adding deterministic header normalization to close the dirty-header failure "
"mode, rejecting any generated specification that references a non-existent column, and extending the assistant to parse "
"field sitreps and community alerts into clean georeferenced records with human verification, alongside multi-disease "
"coverage for typhoid, dysentery, malaria, and dengue.")
figph("Figure 13", img="/root/paper_figures_gen/diag_roadmap.png")
caption("Figure 13: Advanced integration and automation roadmap — NCDC/SORMAS API, NASRDA satellite data hub, DHS microdata "
"layer, laboratory-confirmation linkage, AI unstructured parser, and calibration/validation engine, unified through an "
"integration gateway.")

# ================= 9. CONCLUSION =================
H("9. Conclusion",1)
P("This study presents the development and pilot validation of a scalable, Earth Observation-enabled environmental health "
"intelligence hub for cholera surveillance in Nigeria, together with a design for national scale-up. Its contribution is "
"reported across four clearly separated tiers of evidence, and the boundary between them is the paper's central "
"methodological commitment.")
P("At the national tier, officially reported NCDC figures—led by the 111,062-case, 3,604-death 2021 epidemic—establish "
"burden context at state resolution as published. At the sentinel tier, a Cross River line-list of 74 suspected cases and "
"4 deaths across four LGAs, of which one was culture-confirmed for Vibrio cholerae, was used to validate ingestion, record "
"normalization, spatial joining and risk rendering. At the environmental tier, intersecting a 2.65-million-polygon dated "
"flood archive with GRID3 boundaries produced observed exposure measures for the pilot LGAs that satisfy a physical "
"validity constraint; the resulting exposure–burden relationship across four units is inverse, is reported as observed, "
"and is explicitly not interpreted as evidence about the environmental hypothesis. At the capability tier, a benchmark of "
"72 live invocations established that the platform's AI copilot invents column names in 97.2% of unconstrained trials, and "
"that requiring it to inspect the data before describing it eliminates that failure entirely while raising correct "
"case-field binding from 33.3% to 91.7%.")
P("What the study does not establish is stated with equal clarity. It does not construct a 774-LGA national panel; it does "
"not calibrate or prospectively validate the risk score; it does not demonstrate forecasting skill; it does not evidence "
"live institutional integration; and it does not generalize its four-LGA pilot to the federation. A candidate national "
"panel was inspected, found irreconcilable with official reporting, and excluded in full, and no synthetic epidemiological "
"value appears anywhere in this work.")
P("Framed honestly against that evidence base, the hub represents a credible technical foundation for transitioning "
"Nigerian cholera surveillance from reactive response toward proactive, environmentally contextualized, and sovereign "
"public-health intelligence. The measured copilot failure mode and its architectural mitigation are offered as a "
"transferable contribution: AI components entering public-health practice should be published with a failure rate and a "
"stated mitigation, in the manner of a diagnostic assay, rather than demonstrated with a single successful screenshot.")

# ================= 10. REFERENCES =================
H("10. References",1)
refs=[
"Adegoke, B. O., Odugbose, T., & Adeyemi, C. (2024). Data analytics for predicting disease outbreaks: A review of models and tools. International Journal of Life Science Research Updates, 2(2), 1-9.",
"Adesina, M. A., Adedeji, R. T., Oladipupo, I. R., Olufadewa, I. I., Oladele, R. I., Olufadewa, T. A., ... & Olansile, A. K. (2026). Cholera amid the climate crisis: a systematic review of flooding-driven health risks in West Africa. Discover Public Health, 23(1), 497.",
"Akanda, A. S., Jutla, A. S., Alam, M., De Magny, G. C., Siddique, A. K., Sack, R. B., ... & Islam, S. (2011). Hydroclimatic influences on seasonal and spatial cholera transmission cycles: implications for public health intervention in the Bengal Delta. Water Resources Research, 47(3).",
"Akingbola, A., Abiodun, A., Ojo, O., Jessica, O. U., Alao, U. H., Owolabi, A. O., & Chuku, J. (2025). Cholera outbreak in Nigeria: history, review of socioeconomic and meteorological drivers, diagnostic challenges, and artificial intelligence integration. Global Health, Epidemiology and Genomics, 2025(1), 8898076.",
"Albertini, C., Gioia, A., Iacobellis, V., & Manfreda, S. (2022). Detection of surface water and floods with multispectral satellites. Remote Sensing, 14(23), 6005.",
"Anderson, K., Ryan, B., Sonntag, W., Kavvada, A., & Friedl, L. (2017). Earth observation in service of the 2030 Agenda for Sustainable Development. Geo-spatial Information Science, 20(2), 77-96.",
"Armando, C. J., Rocklöv, J., Sidat, M., Tozan, Y., Mavume, A. F., Bunker, A., & Sewe, M. O. (2024). Spatial-temporal analysis of climate and socioeconomic conditions on cholera incidence in Mozambique from 2000 to 2018. BMJ Open, 14(8), e082503.",
"Atobatele, O. K., Ajayi, O. O., Hungbo, A. Q., & Adeyemi, C. (2019). Leveraging public health informatics to strengthen monitoring and evaluation of global health intervention. IRE Journals, 2(7), 174-193.",
"Avtar, R., Komolafe, A. A., Kouser, A., Singh, D., Yunus, A. P., Dou, J., ... & Kumar, P. (2020). Assessing sustainable development prospects through remote sensing: A review. Remote Sensing Applications: Society and Environment, 20, 100402.",
"Bhaga, T. D., Dube, T., Shekede, M. D., & Shoko, C. (2023). Investigating the effectiveness of Landsat-8 OLI and Sentinel-2 MSI satellite data in monitoring the effects of drought on surface water resources. Remote Sensing Applications: Society and Environment, 32, 101037.",
"Bhunia, G. S., & Shit, P. K. (2021). GeoComputation and Disease Ecology. In GeoComputation and Public Health: A Spatial Approach (pp. 151-220). Springer.",
"Bose, I., Hadida, G., Green, R., Murray, K. A., Part, C., & Kovats, S. (2026). Rainfall and water-related diseases, malnutrition and mortality in Low- and Middle-Income Countries: a systematic review. Heliyon, 12(1).",
"Brumfield, K. D., Usmani, M., Long, D. M., Lupari, H. A., Pope, R. K., Jutla, A. S., ... & Colwell, R. R. (2025). Climate change and Vibrio: Environmental determinants for predictive risk assessment. PNAS, 122(33), e2420423122.",
"Campbell, A. M., Racault, M. F., Goult, S., & Laurenson, A. (2020). Cholera risk: a machine learning approach applied to essential climate variables. International Journal of Environmental Research and Public Health, 17(24), 9378.",
"Chowdhury, F., Ross, A. G., Islam, M. T., McMillan, N. A. J., & Qadri, F. (2022). Diagnosis, Management, and Future Control of Cholera. Clinical Microbiology Reviews, 35(3), e0021121.",
"Chowdhury, A. H., & Rahman, M. S. (2025). Machine learning and spatio-temporal analysis of meteorological factors on waterborne diseases in Bangladesh. PLoS Neglected Tropical Diseases, 19(1), e0012800.",
"Christaki, E., Dimitriou, P., Pantavou, K., & Nikolopoulos, G. K. (2020). The impact of climate change on cholera: A review on the global status and future challenges. Atmosphere, 11(5), 449.",
"Cord, A. F., Brauman, K. A., Chaplin-Kramer, R., Huth, A., Ziv, G., & Seppelt, R. (2017). Priorities to advance monitoring of ecosystem services using earth observation. Trends in Ecology & Evolution, 32(6), 416-428.",
"da Silva Junior, U. J., da Penha Pacheco, A., Ruiz-Armenteros, A. M., et al. (2025). A methodological approach to flood dynamics based on satellite-derived spectral indices and altimetric forecast models. Science of the Total Environment, 1003, 180686.",
"Escobar, L. E., Ryan, S. J., Stewart-Ibarra, A. M., et al. (2015). A global map of suitability for coastal Vibrio cholerae under current and future climate conditions. Acta Tropica, 149, 202-211.",
"Ganesan, D., Gupta, S. S., & Legros, D. (2020). Cholera surveillance and estimation of burden of cholera. Vaccine, 38, 13-17.",
"Geng, Q., Yao, X., Sun, J., Jia, S., Xu, F., Zhang, L., ... & Li, G. (2026). Open-source GIS software and tools: a systematic review. Big Earth Data, 1-48.",
"Gopo, L., Bere, T., & Murisa, M. R. (2026). Climatic and socio-environmental drivers of cholera epidemics. PLOS Climate, 5(4), e0000840.",
"Gorelick, N., Hancher, M., Dixon, M., Ilyushchenko, S., Thau, D., & Moore, R. (2017). Google Earth Engine: Planetary-scale geospatial analysis for everyone. Remote Sensing of Environment, 202, 18-27.",
"Hashizume, M., Armstrong, B., Hajat, S., Wagatsuma, Y., Faruque, A. S. G., Hayashi, T., & Sack, D. A. (2008). The effect of rainfall on the incidence of cholera in Bangladesh. Epidemiology, 19(1), 103-110.",
"Huffman, G. J., Stocker, E. F., Bolvin, D. T., Nelkin, E. J., & Tan, J. (2020). GPM IMERG Final Precipitation L3 Half-Hourly 0.1° x 0.1°. NASA GES DISC.",
"Kamalrathne, T., Amaratunga, D., Haigh, R., & Kodituwakku, L. (2023). Need for effective detection and early warnings for epidemic and pandemic preparedness planning. International Journal of Disaster Risk Reduction, 92, 103724.",
"Kanagaraj, J. K., & Vijayan, T. B. (2024). Spatio-temporal correlation analysis of environmental and climatic determinants with various infectious diseases in Tamil Nadu and West Bengal. Environmental Engineering & Management Journal, 23(9).",
"Khachoo, Y. H., Cutugno, M., Robustelli, U., & Pugliano, G. (2026). Google Earth Engine since 2022: A structured bibliometric review of GeoAI-driven trends and applications. Sustainability, 18(12), 6241.",
"Kuna, A., & Gajewski, M. (2017). Cholera—the new strike of an old foe. International Maritime Health, 68(3), 163-167.",
"Li, X., Zhou, Z., Jia, H., Li, Z., Yang, Z., Cai, Z., ... & Shi, X. (2026). Mechanisms of accumulation–transport–discharge and source apportionment of combined sewer overflow pollution. Water, 18(5), 573.",
"Lloyd, C. T., Sturrock, H. J. W., Leasure, D. R., Jochem, W. C., Lázár, A. N., & Tatem, A. J. (2020). Using GIS and machine learning to classify residential status of urban buildings in low- and middle-income settings. Remote Sensing, 12(23), 3847.",
"Madubueze, M. H., Mbanefo, O. D., Anekwe, J. K., Nwadiogbu, N. M., & Egberi, A. E. (2025). Climate variability and waterborne disease burden in Nigeria. International Journal of Interdisciplinary Studies and Innovation, 4(1), 85-90.",
"McFeeters, S. K. (1996). The use of the Normalized Difference Water Index (NDWI) in the delineation of open water features. International Journal of Remote Sensing, 17(7), 1425-1432.",
"Meckawy, R., Stuckler, D., Mehta, A., Al-Ahdal, T., & Doebbeling, B. N. (2022). Effectiveness of early warning systems in the detection of infectious diseases outbreaks: a systematic review. BMC Public Health, 22(1), 2216.",
"Mikaberidze, A., et al. (2025). Opportunities and challenges in combining optical sensing and epidemiological modeling. Phytopathology, 115(10), 1260-1285.",
"Mobasheri, M. R. (2022). Remote Sensing and Computational Epidemiology. In COVID-19 Pandemic, Geospatial Information, and Community Resilience (pp. 55-68). CRC Press.",
"Morse, S. S., Mazet, J. A., Woolhouse, M., Parrish, C. R., Carroll, D., Karesh, W. B., ... & Daszak, P. (2012). Prediction and prevention of the next pandemic zoonosis. The Lancet, 380(9857), 1956-1965.",
"National Population Commission & ICF. (2024). Nigeria Demographic and Health Survey 2023-24: Key Indicators Report. Abuja & Rockville.",
"Nigeria Centre for Disease Control and Prevention (NCDC). (2022). Cholera Situation Report, Epi Week 52, 2021. Abuja: NCDC.",
"Nigeria Centre for Disease Control and Prevention (NCDC). (2024). Cholera Situation Report, Epi Week 22, 2024. Abuja: NCDC.",
"Nigeria Centre for Disease Control and Prevention (NCDC). (2025). An update of cholera outbreak in Nigeria. Federal Ministry of Health.",
"Olagunju, A. T., Denga, V. S., Madakadze, C., & Osuagwu, F. (2025). Climate change and health issues: an Afrocentric perspective. In Health and Climate Change (pp. 251-265). Academic Press.",
"Peprah, M. S., Moomen, A. W., & Sey, J. M. (2026). Leveraging Digital Earth for sustainable development in Africa: a systematic review and meta-analysis. Arabian Journal of Geosciences, 19(5), 99.",
"Pezanowski, S., Koua, E. L., Okeibunor, J. C., & Gueye, A. S. (2024). Predictors of disease outbreaks at continental scale in the African region. Digital Health, 10, 20552076241278939.",
"Reyburn, R., Kim, D. R., Emch, M., Khatib, A., Von Seidlein, L., & Ali, M. (2011). Climate variability and the outbreaks of cholera in Zanzibar, East Africa: a time series analysis. American Journal of Tropical Medicine and Hygiene, 84(6), 862.",
"Schets, F. M., Pol-Hofstad, I. E., van den Berg, H. H., & Schijven, J. F. (2025). Climate change-related temperature impact on human health risks of Vibrio species in bathing and surface water. Microorganisms, 13(8), 1893.",
"Singh, K., Kumar, S., Brumfield, K. D., Deliz, K., Colwell, R. R., Jutla, A., & Usmani, M. (2025). Remote Sensing to Predict Climate-Sensitive Pathogens: a Review. Authorea Preprints.",
"Singh, S., Sharma, P., Pal, N., Sarma, D. K., Tiwari, R., & Kumar, M. (2024). Holistic one health surveillance framework. ACS Infectious Diseases, 10(3), 808-826.",
"Song, Y., & Wu, P. (2021). Earth observation for sustainable infrastructure: A review. Remote Sensing, 13(8), 1528.",
"Tucker, C. J. (1979). Red and photographic infrared linear combinations for monitoring vegetation. Remote Sensing of Environment, 8(2), 127-150.",
"United Nations Children's Fund (UNICEF). (2025). Nigeria Humanitarian Flash Update (Cholera Outbreak), 1, 1-6.",
"Usmani, M., Brumfield, K. D., Magers, B. M., et al. (2023). Combating cholera by building predictive capabilities for pathogenic Vibrio cholerae in Yemen. Scientific Reports, 13(1), 2255.",
"Verma, R., & Kotwal, M. (2025). Geographic information systems: a review of its evolution, challenges, and future trends. International Journal for Multidisciplinary Research.",
"Warekuromor, T. (2026). Development and Pilot Validation of the Nigeria EO-enabled Environmental Health Intelligence Hub for Cholera Surveillance [Conference presentation]. GEO Health Community of Practice, NASRDA.",
"Wright, A. K. A., Ezugwu, C. I., Iregbu, J. K., et al. (2025). Climate change and emerging infectious diseases: a global review. Epidemiology and Health Data Insights, 1(3), ehdi009.",
"Yaghobi, S., Daneshi, A., Faramarzi, M., Azadi, H., Fathizad, H., & Islami, I. (2026). Drought impacts on vegetation cover in Western Iran. International Journal of Environmental Science and Technology, 23(1), 35.",
"Yu, H. (2025). Climate change unveils hidden microbial dangers. Environmental Science and Ecotechnology, 24, 100544.",
"Zhao, Q., Yu, L., Du, Z., Peng, D., Hao, P., Zhang, Y., & Gong, P. (2022). An overview of the applications of earth observation satellite data. Remote Sensing, 14(8), 1863.",
]
for r in sorted(refs):
    p=doc.add_paragraph(); run=p.add_run(r); run.font.size=Pt(9)
    p.paragraph_format.left_indent=Pt(18); p.paragraph_format.first_line_indent=Pt(-18)

# ================= DECLARATIONS =================
H("Declarations",1)
P("Ethics approval. This study used publicly available, officially reported aggregate epidemiological data and an "
"anonymized outbreak line-list supplied through official channels. No personally identifiable information was accessed "
"and no human subjects research was undertaken.",size=10)
P("Data availability. Nationally reported cholera figures are available in the published NCDC situation reports cited in "
"Section 10. The Cross River 2021 line-list contains outbreak response data and is available from the Cross River State "
"Ministry of Health subject to their data-sharing governance. GRID3 administrative boundaries and the FMOH health-facility "
"registry are publicly available from their respective providers.",size=10)
P("Reproducibility of the copilot benchmark. The benchmark reported in Section 6 is fully reproducible from the pilot "
"dataset. The twelve schema variants are generated deterministically from the observed line-list; the harness records "
"every trial with its variant, condition, repetition, emitted specification, referenced columns, latency and outcome; "
"infrastructure failures are detected, retried and excluded from scoring; and every trial is asserted to have used a live "
"model invocation rather than the offline fallback path. Benchmark code and the complete 72-trial result set are available "
"from the corresponding author on request.",size=10)
P("Author contributions. W.T. conceived the study, leads the Mission Planning programme under which it was undertaken, and "
"supervised the work. A.M.O. contributed to the Earth Observation methodology and institutional framing. Y.T.U. "
"implemented the platform, performed the data engineering, spatial analysis and copilot benchmark, and prepared the "
"figures. All authors reviewed and approved the manuscript.",size=10)
P("Conflicts of interest. The authors declare no competing interests.",size=10)
P("Funding. This work was undertaken within the normal programme of the National Space Research and Development Agency "
"(NASRDA). No external funding was received.",size=10)
P("Use of AI tools. An AI assistant integrated into the platform is a subject of study in Section 6 and its behaviour is "
"reported as measured experimental data. AI-based tooling was additionally used to support data engineering and analysis "
"scripting; all numerical results were computed programmatically from source data and verified by the authors, and no "
"epidemiological value reported in this paper was generated by a language model.",size=10)

doc.save('/root/CHOLERA_PAPER_V6_Revised.docx')
print("V6 saved. paragraphs:", len(doc.paragraphs), "| tables:", len(doc.tables), "| references:", len(refs))
