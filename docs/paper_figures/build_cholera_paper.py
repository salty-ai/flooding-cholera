# -*- coding: utf-8 -*-
"""Build honest real-data cholera manuscript V5 from HoD original base.
(2) real NCDC data + (3) platform/methods framing. No synthetic numbers."""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()
normal = doc.styles['Normal']; normal.font.name='Times New Roman'; normal.font.size=Pt(11)
CENTER=WD_ALIGN_PARAGRAPH.CENTER; JUST=WD_ALIGN_PARAGRAPH.JUSTIFY

def H(t,l=1):
    p=doc.add_heading(t,level=l)
    for r in p.runs: r.font.color.rgb=RGBColor(0,0,0)
    return p
def P(t,italic=False,bold=False,align=JUST,size=11):
    p=doc.add_paragraph(); r=p.add_run(t); r.italic=italic; r.bold=bold; r.font.size=Pt(size)
    p.alignment=align; return p
def caption(t):
    p=doc.add_paragraph(); r=p.add_run(t); r.italic=True; r.font.size=Pt(9)
    p.alignment=CENTER; return p
# map figure-slot -> real image file
import os
from docx.shared import Inches
_IMG={
 "Figure 1":"/root/paper_figures_gen/diag_architecture.png",
 "Figure 2":"/root/paper_figures_gen/diag_pipeline.png",
 "Figure 3":"/root/paper_figures_app/app_dashboard.png",
 "Figure 4":"/root/paper_figures_app/app_map.png",
 "Figure 5":"/root/paper_figures_gen/fig_crossriver_pilot.png",
 "Figure 6":"/root/paper_figures_app/app_facilities.png",
 "Figure 7":"/root/paper_figures_gen/fig_national_burden.png",
 "Figure 8":"/root/paper_figures_gen/fig_lag_signal.png",
 "Figure 9":"/root/paper_figures_app/app_alerts.png",
 "Figure 10":"/root/paper_figures_app/app_agent.png",
 "Figure 11":"/root/paper_figures_gen/diag_roadmap.png",
}
_used=set()
def figph(t, img=None, width=6.2):
    """Embed a real image centered. img overrides keyword mapping."""
    path=img
    if path is None:
        for k,v in _IMG.items():
            if t.startswith(k): path=v; break
    p=doc.add_paragraph(); p.alignment=CENTER
    if path and os.path.exists(path):
        p.add_run().add_picture(path, width=Inches(width))
    else:
        r=p.add_run("[ "+t+" ]"); r.italic=True; r.font.size=Pt(9); r.font.color.rgb=RGBColor(0x88,0x88,0x88)
    return p
def fig(img, width=6.2):
    p=doc.add_paragraph(); p.alignment=CENTER
    if os.path.exists(img): p.add_run().add_picture(img, width=Inches(width))
    return p
def table(headers,rows):
    t=doc.add_table(rows=1,cols=len(headers)); t.style='Light Grid Accent 1'; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; c.text=''; r=c.paragraphs[0].add_run(h); r.bold=True; r.font.size=Pt(9)
    for row in rows:
        cells=t.add_row().cells
        for i,v in enumerate(row):
            cells[i].text=''; r=cells[i].paragraphs[0].add_run(str(v)); r.font.size=Pt(9)
    return t

# ================= TITLE =================
tp=doc.add_paragraph(); tr=tp.add_run("Development and Pilot Validation of a Scalable Earth Observation-Enabled "
"Environmental Health Intelligence Hub for Cholera Surveillance in Nigeria")
tr.bold=True; tr.font.size=Pt(15); tp.alignment=CENTER
P("Warekuromor Tubolayefa, Adepoju Matthew Olumide, Yakubu Tanimu Umar",align=CENTER,bold=True)
P("Head, Mission Planning, Department of Mission Planning and Satellite Data Management.",align=CENTER,italic=True,size=10)
P("National Space Research and Development Agency (NASRDA), Abuja-Nigeria",align=CENTER,italic=True,size=10)

# ================= ABSTRACT =================
H("Abstract",1)
P("This study presents the development, pilot demonstration, and national scale-up design of an "
"Earth Observation (EO)-enabled environmental health surveillance and intelligence hub for cholera in Nigeria. "
"The platform harmonizes epidemiological records with satellite-derived environmental indicators accessed through "
"Google Earth Engine (GEE)—including NASA Global Precipitation Measurement (GPM-IMERG) rainfall, the Normalized "
"Difference Water Index (NDWI) from Sentinel-2, the Normalized Difference Vegetation Index (NDVI) from Landsat, and "
"Synthetic Aperture Radar (SAR) flood extent from Sentinel-1—to characterize environmental drivers of waterborne "
"disease across space and time and to compute an interpretable, LGA-level composite risk score within a geostatistical "
"multi-criteria framework.")
P("A 30-day pilot demonstration for Cross River State ingested a culture-referenced cholera line-list of 74 confirmed "
"cases and 4 deaths (case fatality rate 5.4%) reported across four Local Government Areas (Yakurr, Biase, Calabar "
"Municipal and Bakassi) during 2021, and was used to validate dashboard ingestion, spatial joins, and the automated "
"risk-rendering workflow rather than to establish epidemiological effect sizes. For national context, the hub compiled "
"officially reported cholera figures from Nigeria Centre for Disease Control and Prevention (NCDC) situation reports, "
"which document 111,062 suspected cases and 3,604 deaths across 33 states and the FCT in 2021—the largest recent "
"national epidemic—followed by substantially lower burdens in 2022–2025. National epidemiological reporting is "
"published at the state level; sub-national LGA-level and ward-level resolution is a design objective of the platform, "
"realized empirically only where line-list data are available.")
P("Cross-correlation between satellite-detected surface-water anomalies and reported cholera activity is provided as an "
"exploratory decision-support signal, consistent with the one-to-two-month environmental lag reported in the literature, "
"and is not presented as a validated forecast or as evidence of causation. The results demonstrate the technical "
"feasibility and operational value of an EO-integrated, multi-agency situational-awareness platform, and set out the "
"data-governance and validation steps required to move from retrospective decision support toward prospective early warning.")
P("Keywords: Earth Observation, Cholera Surveillance, Hydro-Epidemiology, Satellite Remote Sensing, Geostatistical Risk "
"Scoring, Flood Risk Mapping, Decision Support Systems, Nigeria.",italic=True,size=10)

# ================= INTRODUCTION (HoD verbatim) =================
H("1. Introduction",1)
P("Public health surveillance and response is often complicated by the complex interaction between hydrometeorological "
"variables and epidemiological outcomes. Traditional public health surveillance is mostly fragmented and lacks the spatial "
"granularity, as well as the real-time environmental data required to holistically understand the triggers and impact for "
"both proactive and retrospective interventions (Singh et al., 2024; Atobatele et al., 2019). In Nigeria, this challenge is "
"magnified by precarious water insecurity, rapid urban growth, and fragile infrastructure, which collectively intensify the "
"threat of waterborne pathogens.")
P("Environment and climate-induced shifts, such as erratic rainfall, increased temperatures, and amplified flood and drought "
"cycles are altering pathogen ecology and transmission, presenting new challenges to the epidemiology of waterborne diseases "
"like cholera (Akingbola et al., 2025; Madubueze et al., 2025). Cholera (Vibrio cholerae) remains a critical public health "
"challenge in Nigeria and across sub-Saharan Africa, exhibiting severe seasonal spikes driven by localized flooding, poor "
"drainage systems and compromised water, sanitation, and hygiene (WASH) infrastructure. Cholera is primarily characterized by "
"acute watery diarrhoea and extreme dehydration (Madumelu et al., 2026; Kuna and Gajewski, 2017). The effects of climate "
"change consistently alter the epidemiology of waterborne diseases through erratic temperature and precipitation patterns, as "
"well as the increased frequency and intensity of extreme weather events such as flooding and drought. These impacts are "
"consistently escalating the environmental conditions that favour the growth of Vibrio cholerae, thereby amplifying public "
"health risks and broadening transmission windows in vulnerable regions (Yu, 2025; Christaki et al., 2020). Rising temperature "
"patterns have been shown to enable the geographic expansion of cholera-endemic conditions as well as the growth of the "
"bacteria, enabling the pathogen to colonize new areas. Meanwhile, erratic rainfall events and their corresponding impacts, "
"including flash and intense flooding, promote the contamination of water supplies through sewer overflow and latrine "
"inundation, directly and indirectly driving spikes in cholera transmission.")
P("Surveillance and monitoring systems are essential for the early detection of outbreaks within highly vulnerable areas. Early "
"intervention through the provision of clean water and oral cholera vaccines, coupled with improved health services, results in "
"a significant potential reduction in the impact of cholera outbreaks (Mikaberidze, 2025; Chowdhury et al., 2022; Meckawy et "
"al., 2022).")
P("This paper seeks to address this critical operational gap by presenting the conceptualization and pilot deployment of an "
"innovative and bespoke Earth Observation (EO)-enabled Environmental Health Surveillance and Intelligence Hub. The platform is "
"designed to integrate spatially referenced climate and environmental hazard data with epidemiological and infrastructure data "
"layers to provide a unified platform for multi-agency situational awareness. To this end, it transforms the surveillance "
"landscape, replacing irregular, reactionary monitoring with a robust geostatistical decision-support system. This system not "
"only identifies the environmental conditions that precede outbreaks but also delivers granular vulnerability assessments "
"disaggregated to the lowest administrative level for which reliable data exist, enabling resource allocation and response "
"strategies to be tailored to the most at-risk communities within Nigeria, leaving no one behind.")

# ================= 2. EO FOR DISEASE SURVEILLANCE (HoD verbatim lit review) =================
H("2. The Use of Earth Observation (EO) for Disease Surveillance",1)
H("2.1 The Capabilities and Applications of Earth Observation (EO)",2)
P("Earth Observation (EO) refers to the systematic gathering of environmental data about the Earth's physical, chemical, and "
"biological systems through remote sensing technologies, including satellite, aerial, and unmanned aerial vehicle (UAV) "
"platforms (Zhao et al., 2022; Song and Wu, 2021; Cord et al., 2017). These sensors capture a wide spectrum of biophysical "
"variables such as land and water surface temperature, precipitation, soil moisture, vegetation health, atmospheric "
"composition, water quality, land cover, and land use dynamics, without requiring physical contact with the surface. Data are "
"typically stored in raster format, where each pixel represents a geolocated measurement, enabling spatially continuous "
"monitoring across vast regions.")
P("Over the years, EO has evolved from a specialized geospatial scientific tool into a cross-cutting enabler of decision-making "
"across diverse sectors, contributing significantly to sustainable development. It is notably a critical component of "
"interdisciplinary research, with applications spanning health, environmental monitoring, disease surveillance, biodiversity "
"conservation, and habitat characterization (Peprah et al., 2026; Mobasheri, 2022; Avtar et al., 2020; Anderson et al., 2017). "
"This interdisciplinary approach has seen significant successes in the detection, evaluation, and mapping of factors affecting "
"public health, enabling environmental and public health administrations to work collaboratively, improve public health "
"outcomes, disaster risk reduction and efficient resource management. Rather than operating in isolation, EO is most "
"effectively harnessed through cross-sectoral integration, linking environmental monitoring with human health, food security, "
"climate adaptation, and other interrelated domains. For example, vegetation indices and soil moisture data inform drought "
"early warnings, sea surface temperature irregularities help track environmental stress on coastal ecosystems, and "
"satellite-based flood mapping has enabled rapid humanitarian responses during flood incidents.")
P("Critically, EO also plays a central role in monitoring Water, Sanitation, and Hygiene (WASH) infrastructure risks, especially "
"in low-resource settings where conventional monitoring is sparse or absent. By detecting environmental precursors such as "
"flooding, stagnant water accumulation, and temperature shifts, EO provides scalable, near real-time data streams that can "
"inform public health interventions before outbreaks occur, as well as after (Pezanowski et al., 2024; Mobasheri, 2022; Peprah "
"et al., 2026). Thus, EO functions as an environmental health infrastructure that offers a continuous, transboundary data "
"pipeline connecting ecological change to human well-being. Many scientists and resource managers have already recognized the "
"importance of adopting predictive approaches for epidemics and pandemics before their occurrences (Adegoke et al., 2024; "
"Kamalrathne et al., 2023; Morse et al., 2012), and in parallel, the use of EO and GIS approaches has noticeably increased.")

H("2.2 Key Environmental Indicators for Cholera Risk Mapping",2)
P("Various studies have demonstrated that satellite-derived environmental indices are critical for modelling the ecological "
"niche of Vibrio cholerae and identifying high-risk areas before outbreaks occur (Li et al., 2026; Adesina et al., 2026; "
"Kanagaraj and Vijayan, 2024; Madubueze et al., 2025; Akingbola et al., 2025; Bhunia and Shit, 2021; Christaki et al., 2020; "
"Ganesan et al., 2020; Campbell et al., 2020; Escobar et al., 2015). These indices serve as representative measurements for "
"hydrological, thermal, and land surface conditions that promote pathogen persistence and transmission. The selection of "
"indicators in this study is informed by extensive literature, with variables chosen based on their demonstrated predictive "
"value, operational availability, and relevance to Nigeria's diverse ecological zones, from the flood-prone Niger Delta to the "
"drought-vulnerable Sahel.")
P("Precipitation and Hydrological Extremities. Rainfall volume and intensity are primary drivers of surface runoff, sewer "
"overflow, and contamination of drinking water sources (Li et al., 2026; Adesina et al., 2026; Madubueze et al., 2025; "
"Christaki et al., 2020). Previous studies have established that precipitation anomalies, particularly extreme rainfall events, "
"are strongly associated with cholera outbreaks in both coastal and inland settings. Critically, the influence of precipitation "
"on cholera incidence often operates with a lag of one to two months, as heavy rains first accumulate in the environment before "
"contaminating water sources (Armando et al., 2024; Campbell et al., 2020; Reyburn et al., 2011). This lag effect underscores "
"the importance of continuous, near-real-time precipitation monitoring for early warning systems.",bold=False)
P("Surface Water Extent and Flood Persistence. The presence of standing, stagnant surface water bodies is a critical "
"environmental determinant of cholera risk. Vibrio cholerae proliferates in warm, brackish, and nutrient-rich waters, often in "
"association with aquatic zooplankton and phytoplankton (Madumelu et al., 2026; Kuna and Gajewski, 2017). Satellite-derived "
"indices such as the Normalized Difference Water Index (NDWI) have been widely validated for detecting open water features and "
"monitoring flood extent (da Silva Junior et al., 2025; Bhaga et al., 2023; Albertini et al., 2022; McFeeters, 1996). The index "
"is calculated as:")
P("NDWI = (Green − NIR) / (Green + NIR)",align=CENTER,italic=True)
P("Values greater than zero (NDWI > 0) indicate the presence of open water features. In the context of cholera risk mapping, "
"NDWI serves as a proxy for surface water availability and flood persistence—conditions that facilitate human exposure to "
"contaminated water.")
P("Soil Moisture and Vegetation Health. Soil moisture serves as a proxy for environmental humidity and residual moisture "
"retention, conditions that prolong the survival of pathogens (Chowdhury and Rahman, 2025; Campbell et al., 2020). The "
"Normalized Difference Vegetation Index (NDVI), originally developed by Tucker (1979), serves as an indirect optical proxy, "
"assessing soil moisture availability by tracking variations in vegetation health, canopy cover, and biomass density (Yaghobi "
"et al., 2026; Gopo et al., 2026; Singh et al., 2025). The index is calculated as:")
P("NDVI = (NIR − Red) / (NIR + Red)",align=CENTER,italic=True)
P("Higher NDVI values reflect conditions of high humidity and moisture retention that provide optimal breeding grounds for "
"pathogens, increasing their persistence and possible transmission.")
P("Increased Temperature and Climate Change. Laboratory and field studies have demonstrated that Vibrio cholerae growth rates "
"increase with temperature, with optimal conditions occurring in warm, brackish waters (Schets et al., 2025; Brumfield et al., "
"2025). Epidemiological studies have shown that a one-degree increase in temperature can double the risk of a cholera outbreak "
"during rainy seasons, particularly in areas without safe water or sanitation (Gopo et al., 2026; Olagunju et al., 2025; "
"Campbell et al., 2020).")
P("Lag Effects and Temporal Dynamics. A critical consideration in environmental cholera modelling is the presence of lag "
"effects between environmental anomalies and subsequent outbreaks. Campbell et al. (2020) found that chlorophyll-a "
"concentration was most predictive when lagged by one to two months, while land surface temperature was most predictive in the "
"current month. Similarly, Hashizume et al. (2008) found that rainfall effects on cholera incidence peaked at a lag of one to "
"two months, reflecting the time required for environmental contamination to translate into human infection. These findings "
"inform the design of early warning systems, which must account for temporal dynamics to provide actionable lead time.")
P("Machine Learning and Geostatistical Approaches. Several studies have explored machine learning techniques, particularly "
"Random Forest classifiers, for cholera risk prediction using remotely-sensed environmental variables (Singh et al., 2025; "
"Usmani et al., 2023; Campbell et al., 2020). Campbell et al. (2020) achieved a sensitivity of 0.895 for cholera outbreak "
"prediction in coastal India, and Usmani et al. (2023) employed near-real-time remotely sensed data to achieve 92% spatial "
"correspondence between predicted regions and areas of reported cases. However, machine learning approaches require extensive "
"training datasets and may face challenges related to data imbalance, overfitting, and interpretability. The present study "
"adopts a geostatistical risk-scoring framework that complements machine learning by providing interpretable, threshold-based "
"risk indicators that can be operationalized within existing public health workflows.")

H("2.3 The Persistent Gap in Infectious Disease Surveillance",2)
P("Despite these advances, the operational use of EO in infectious disease surveillance—particularly for climate-sensitive, "
"waterborne pathogens like Vibrio cholerae—remains limited and fragmented, especially in high-burden, resource-limited regions "
"(Wright et al., 2025; Singh et al., 2025; Akanda et al., 2011). While decades of research have established strong associations "
"between environmental drivers and cholera outbreaks, this knowledge has not been systematically translated into public health "
"practice. Singh et al. (2025) synthesized 1,137 peer-reviewed studies on remote sensing applications in infectious disease "
"surveillance and found that predictive approaches are significantly more effective than reactive responses. However, they "
"identified persistent challenges that limit operational use, particularly the difficulty of obtaining field-collected "
"environmental data aligned with microbiological measurements, and inconsistencies in epidemiological reporting. These barriers "
"are especially pronounced in settings such as Nigeria, where surveillance systems are often fragmented, data availability is "
"low, and cross-sectoral coordination remains weak.")
P("Yet, despite this well-documented environmental epidemiology, most public health systems in West Africa continue to rely on "
"reactive, case-based surveillance—responding after outbreaks are detected rather than anticipating them. This reactive posture "
"stems from a critical disconnect between the geospatial and health sectors: while EO data can detect environmental risk weeks "
"in advance, these signals rarely reach decision-makers in time or in a usable format.")

H("2.4 The Need for Scalable Cloud-Based Geospatial Systems",2)
P("The use of traditional Geographic Information Systems (GIS) in the analysis of spatial health data is not novel. However, "
"traditional desktop GIS continues to fall short compared to modern EO workflows due to significant limitations in computational "
"capacity, data storage, and real-time processing ability (Geng et al., 2026; Verma and Kotwal, 2025; Lloyd et al., 2020). The "
"volume of satellite data, often measured in petabytes, exceeds the processing capabilities of most local institutions, "
"particularly in low-resource settings.")
P("Cloud-based geospatial platforms offer a transformative alternative. Platforms such as Google Earth Engine (GEE), Amazon Web "
"Services, and the Microsoft Planetary Computer provide instant access to decades of satellite imagery, pre-processed data "
"streams, and massive computing power directly within a web browser (Khachoo et al., 2026; Yu, 2025; Gorelick et al., 2017). "
"When combined with open-source software such as QGIS, SNAP, and web GIS servers, organizations can process, publish, and present "
"spatial data without relying on expensive proprietary software. However, global platforms also raise concerns around data "
"sovereignty, national security, and long-term sustainability, particularly when sensitive health and environmental data are "
"processed or stored abroad. This study advances national data sovereignty by designing an environmental-health surveillance hub "
"that integrates global satellite data with locally sourced observations and incident data, and that is engineered to function "
"even if access to global data is disrupted, by leveraging domestic remote sensing capacity and ground-based monitoring.")

# ================= 3. METHODOLOGY (rewritten, honest + real provenance) =================
H("3. Methodology and System Architecture",1)
H("3.1 Study Setting and Data Architecture",2)
P("This study develops a cholera surveillance system that supports early warning design, retrospective research, and health data "
"reporting. The intended geographic scope covers all 36 states of the Federal Republic of Nigeria and the Federal Capital "
"Territory (FCT), and, where sub-national data are available, the 774 Local Government Areas (LGAs) that constitute "
"Administrative Level 2. The system utilizes publicly available, anonymized epidemiological data from the situation reports and "
"weekly epidemiological reports of the Nigeria Centre for Disease Control and Prevention (NCDC), together with satellite-derived "
"environmental data from the National Space Research and Development Agency (NASRDA) archives, NASA, and the European Space "
"Agency (ESA). The study did not involve human subjects or personally identifiable information.")
P("An important structural characteristic of the national epidemiological data governs the analysis presented here. NCDC cholera "
"situation reports publish suspected cases, deaths, and case fatality ratios primarily at the state level, with LGA-level detail "
"provided only for a subset of high-burden states in a given reporting period. Consequently, this study presents nationally "
"reported figures at the state level as officially published, and treats full 774-LGA and ward-level resolution as a design "
"objective of the platform that is realized empirically only where sub-national line-list data are supplied by state Ministries "
"of Health or the National Cholera Technical Working Group.")
figph("Figure 1")
caption("Figure 1: System architecture — three-tier Earth Observation-enabled cholera surveillance hub, showing raw data "
"sources (Tier 1), backend microservices (Tier 2), the delivery layer (Tier 3), and actionable outputs.")

H("3.2 Data Sources and Provenance",2)
P("To build a georeferenced public health dataset, the hub ingests and harmonizes several data repositories, summarized with "
"their provenance and evidence status in Table 1.")
P("Table 1: Data sources, provenance, and evidence status.",bold=True,size=10)
table(["Data Layer","Source","Resolution / Coverage","Evidence Status"],
[["National cholera burden","NCDC Cholera Situation Reports (Epi Week reports, 2021–2025)","State level; 30–35 states + FCT per year","Observed (official)"],
 ["Cross River pilot line-list","Cross River SMoH / NCDC cholera line-list, 2021 (culture-referenced)","4 LGAs; 74 case records","Observed (line-list)"],
 ["Administrative boundaries","GRID3 Nigeria Admin Level 2","774 LGAs, 36 states + FCT","Observed"],
 ["Precipitation","NASA GPM-IMERG via Google Earth Engine","0.1°, national","Computed on demand"],
 ["Surface water (NDWI)","Sentinel-2 MSI via GEE","10–20 m","Computed on demand"],
 ["Vegetation (NDVI)","Landsat-8/9 via GEE","30 m","Computed on demand"],
 ["Flood extent (SAR)","Sentinel-1 GRD via GEE; NEMA flood datasets","10 m; national event-based","Computed / observed"],
 ["Health facilities","FMOH Nigeria Health Facility Registry","46,146 facility records (source count)","Observed (unvalidated)"],
 ["Historical flood polygons","Groundsource flood archive","2000–2026, 2.6M polygons","Observed"]])
P("Epidemiological registries: nationally reported cholera figures were compiled from NCDC National Cholera Situation Reports. "
"For 2021, NCDC reported 111,062 suspected cases and 3,604 deaths (CFR 3.2%) across 33 states and the FCT (NCDC, 2022). "
"Sub-national pilot data for Cross River State were drawn from a 2021 cholera line-list containing culture-referenced case "
"records with onset dates, wards, settlements, and outcomes.")
P("Geospatial boundaries: administrative boundary shapefiles for Nigeria's second administrative tier (LGA boundaries) were "
"sourced from GRID3. Satellite and climatic data: remote-sensing indices were ingested programmatically via Google Earth Engine, "
"leveraging NASA GPM-IMERG precipitation, Sentinel-2 and Landsat surface hydrology and vegetation dynamics, and Sentinel-1 SAR "
"flood mapping; NASRDA archive data were also ingested.")

H("3.3 Environmental Covariate Processing (Google Earth Engine)",2)
P("Environmental indices were computed within GEE and aggregated to administrative boundaries. Precipitation depth (mm) was "
"derived from GPM-IMERG as a physical driver of surface runoff and contamination overflow. NDWI was computed from Sentinel-2 "
"Green and Near-Infrared bands, with values greater than zero isolating open water, localized pooling, and riverine flood "
"corridors; the flood-mask implementation applies an operational NDWI threshold of 0.3, which is a configurable parameter "
"subject to local calibration rather than a validated universal constant. NDVI was derived from red and near-infrared "
"reflectance as a proxy for canopy moisture and soil saturation. To reflect the one-to-two-month environmental lag reported in "
"the literature (Campbell et al., 2020; Hashizume et al., 2008), the hub aggregates environmental indices over a rolling window "
"preceding each epidemiological reporting period. Where a satellite product is unavailable for a requested location or date, "
"the platform returns an explicit unavailable status rather than substituting synthetic values.")
figph("Figure 2")
caption("Figure 2: End-to-end data processing pipeline — from raw EO satellite grids and NCDC/line-list epidemiological "
"records through preprocessing, spatial join, and geostatistical risk scoring to alerts and decision support.")

H("3.4 Geostatistical Risk Scoring",2)
P("The hub computes a composite cholera vulnerability score for each administrative unit using a normalized weighted "
"multi-criteria decision analysis (MCDA) framework:")
P("R_i = w1·P_i + w2·NDWI_i + w3·NDVI_i + w4·C_hist_i + w5·HF_i",align=CENTER,italic=True)
P("where P_i is the precipitation anomaly, NDWI_i is flood-water persistence, NDVI_i is vegetation/soil-moisture retention, "
"C_hist_i is recent reported case burden, and HF_i is a health-facility proximity index. In the current implementation the "
"weights are fixed, literature-informed heuristic coefficients rather than statistically calibrated parameters; principal "
"component analysis and outcome-based calibration are identified as validation steps in Section 6, not as implemented features. "
"Because the score incorporates recent reported cases as an input, it is an interpretable decision-support signal for current "
"vulnerability and is explicitly not a prospective forecast: any predictive evaluation would require a temporal hold-out design "
"that excludes contemporaneous case counts from the predictors.")

H("3.5 Data Quality, Completeness, and Latency",2)
P("Epidemiological reporting is subject to latency, under-reporting, and variation in diagnostic capacity across states and "
"LGAs. Where sub-national panels are constructed, LGAs with reporting gaps exceeding three consecutive months are flagged, and "
"only units meeting a completeness threshold are retained for longitudinal analysis. Reported figures reflect suspected cases as "
"published by NCDC and should be interpreted as conservative baselines rather than absolute biological counts. These constraints "
"reinforce the value of automated, direct data pipelines to minimize collection latency.")

# ================= 4. RESULTS =================
H("4. Pilot Results and Surveillance Demonstrator",1)
H("4.1 Phase 1: Cross River Sentinel Pilot Demonstration",2)
P("Prior to national compilation, the hub's ingestion and risk-rendering capability was demonstrated using the Cross River 2021 "
"cholera line-list. The line-list comprised 74 culture-referenced cases and 4 deaths (case fatality rate 5.4%) distributed "
"across four LGAs, with onset spanning epidemiological weeks 6–46. The pilot exercised dashboard ingestion, spatial joins to LGA "
"boundaries, and automated risk categorization; it is a technical demonstration of the workflow rather than an epidemiological "
"effect estimate. Table 2 summarizes the pilot line-list.")
P("Table 2: Cross River 2021 sentinel pilot line-list summary.",bold=True,size=10)
table(["Sentinel LGA","State / Zone","Reported Cases","Deaths","Case Fatality Rate"],
[["Yakurr","Cross River (South-South)","53","0","0.0%"],
 ["Biase","Cross River (South-South)","10","1","10.0%"],
 ["Calabar Municipal","Cross River (South-South)","6","0","0.0%"],
 ["Bakassi","Cross River (South-South)","5","3","60.0%"],
 ["Total","Cross River","74","4","5.4%"]])
figph("Figure 3")
caption("Figure 3: Main dashboard interface of the surveillance hub, displaying the national KPI summary cards, "
"date-range controls, alert level, and the geospatial risk context for the selected surveillance window.")
figph("Figure 4")
caption("Figure 4: Interactive 774-LGA choropleth risk map (MapLibre/Leaflet) with health-facility overlay, "
"risk-level legend, and Time-Lapse animation control.")
P("The dashboard and map successfully rendered geospatial risk layers, LGA-level surveillance reports, and correlation "
"analytics from the ingested line-list, confirming the operational readiness of the ingestion and visualization pipeline. "
"Figure 5 shows the pilot case and death distribution across the four Cross River LGAs.")
figph("Figure 5")
caption("Figure 5: Cross River 2021 sentinel pilot — reported cases and deaths by LGA, derived directly from the "
"culture-referenced line-list (Yakurr, Biase, Calabar Municipal, Bakassi; 74 cases, 4 deaths).")
P("The small case volume and restricted geographic footprint of this pilot preclude inferential conclusions; the exercise "
"establishes technical feasibility and interface usability for public health practitioners. The platform also incorporates the "
"FMOH health-facility registry (Figure 6) to support response planning.")
figph("Figure 6")
caption("Figure 6: FMOH Health Facility Registry overlay — 46,146 source facility records across the 36 states and FCT, with "
"functional-status summary. Record-level validation is a governance step and is not established in this study.")

H("4.2 Phase 2: Nationally Reported Cholera Burden (2021–2025)",2)
P("For national context, officially reported cholera figures were compiled from NCDC situation reports (Table 3). The national "
"burden is dominated by the 2021 epidemic—the largest in recent years—followed by markedly lower burdens in 2022–2025, with "
"year-to-year fluctuation driven by rainfall, flooding, displacement, and WASH conditions. These figures are reported by NCDC at "
"the state level; they are presented here as published and are not disaggregated by this study into a synthetic LGA panel.")
P("Table 3: Nationally reported cholera burden, Nigeria (NCDC situation reports).",bold=True,size=10)
table(["Year","Suspected Cases","Deaths","CFR","States Reporting","Source"],
[["2021","111,062","3,604","3.2%","33 states + FCT","NCDC SitRep Epi Wk 52, 2021"],
 ["2022","~23,550","583+","2.5%","33 states, 270 LGAs","NCDC SitRep Wks 44–47, 2022"],
 ["2023","reduced ~85% vs 2022","reduced ~79%","—","variable","NCDC SitRep Wks 48–52, 2023"],
 ["2024","4,809 (to 21 Jul)","156","3.2%","35 states","NCDC SitRep 2024"],
 ["2025","1,307 (to 20 Apr)","34","2.6%","30 states, 98 LGAs","UNICEF/NCDC, 2025"]])
figph("Figure 7")
caption("Figure 7: Nationally reported cholera cases and deaths by year, 2021–2025, compiled from NCDC situation reports. The "
"2021 epidemic (111,062 cases; 3,604 deaths) represents the largest recent national outbreak.")
P("The 2021 outbreak coincided with above-average rainfall across the Guinea Savannah and Sahel zones, consistent with the "
"environmental determinants reviewed in Section 2. Retrospective overlay of NDWI and flood-extent anomalies against reported "
"cholera activity provides an exploratory decision-support signal for the association between surface-water persistence and "
"subsequent case reporting. Consistent with the literature (Campbell et al., 2020; Hashizume et al., 2008), the environmental "
"signal precedes reported case activity by approximately one to two months; this relationship is presented as a decision-support "
"indicator for prioritization and is not a validated forecast or proof of causation.")
figph("Figure 8")
caption("Figure 8: Exploratory flood–cholera temporal association (illustrative decision-support signal). Dependence structure "
"and multiplicity are not corrected, and no causal or predictive claim is made.")

H("4.3 Platform Capabilities",2)
P("Beyond the epidemiological demonstrations, the platform provides a configurable rule-based alert engine (Figure 9), a "
"conversational assistant with model-routing hooks and an assisted no-code data-ingestion workflow (Agent Explorer, Figure 10), "
"and PDF/CSV report export. These are described as implemented software capabilities; their epidemiological performance (alert "
"sensitivity, geocoding accuracy) requires separate evaluation and is not claimed here.")
figph("Figure 9")
caption("Figure 9: Automated early-warning alert engine — severity and status filters, alert rails, and threshold-rule "
"management for case-surge, high-risk-score, and recent-flooding triggers.")
figph("Figure 10")
caption("Figure 10: Agent Explorer — assisted no-code ingestion of custom CSV/Excel surveillance data, with schema detection "
"and adaptive dashboard generation via the Surveillance Copilot.")

# ================= 5. DISCUSSION =================
H("5. Discussion",1)
H("5.1 Operationalizing Intelligence and Multi-Agency Awareness",2)
P("The hub is designed to support after-action review by retrospectively mapping cholera outbreaks against historical "
"environmental data, allowing health officials to identify recurrent high-risk zones and the environmental conditions—such as "
"standing water detected via NDWI—that facilitate pathogen proliferation. In the 2021 national epidemic, retrospective overlay "
"of surface-water anomalies against reported case activity is consistent with an environmental lead time of several weeks; if "
"such signals were flagged operationally in real time, they could support the earlier prepositioning of oral rehydration salts, "
"chlorination supplies, and vaccine stockpiles. This remains a hypothesis for prospective evaluation rather than a demonstrated "
"operational outcome.")
P("A core strength of the platform is the shared operational dashboard, which can provide a single source of truth for disease "
"control agencies, national and state health agencies, and disaster-management teams. By presenting epidemiological and "
"environmental data on a single interface, the hub is intended to bridge the institutional gap between health agencies (e.g. "
"NCDC) and environmental agencies (e.g. NASRDA, NEMA), enabling coordination based on the same spatial intelligence. Realizing "
"this shared awareness in production depends on data-sharing agreements and interoperability arrangements that are governance "
"prerequisites rather than technical outputs of this study.")

H("5.2 Comparison with Existing Systems",2)
P("Table 4 situates the hub relative to existing international and regional climate-health surveillance frameworks across "
"resolution, sensor integration, latency, data sovereignty, and decision support. The hub's distinguishing design goals are "
"sub-national (LGA-level) targeting, multi-sensor integration, sovereign hosting, and an assisted analytics layer.")
P("Table 4: Comparative framework of climate-health intelligence platforms.",bold=True,size=10)
table(["Platform","Spatial Resolution","Sensor Integration","Latency","Data Sovereignty","Decision Support"],
[["UNICEF cholera platform","Global / Admin-1","Precipitation, temperature","Monthly batch","Foreign cloud","Static risk mapping"],
 ["WHO Health-Earth","Regional / continental","Precipitation, SST","Bi-weekly / monthly","International servers","Predictive modelling"],
 ["ECDC EWARS","National / regional","Meteorological grids, land cover","Weekly sitreps","European cloud","Rule-based alerts"],
 ["NASRDA EO-Hub (this study)","Sub-national (LGA) design; state where reported","GPM-IMERG, Sentinel-2 NDWI, Landsat NDVI, Sentinel-1 SAR","On-demand retrieval","Sovereign Nigerian infrastructure + NASRDA archives","Rule-based alerts + assisted analytics"]])

H("5.3 Policy Implications",2)
P("The platform is intended to enable three policy shifts: (1) prepositioning of oral rehydration salts, vaccines, and other "
"countermeasures in high-risk LGAs ahead of peak transmission; (2) coordinated inter-agency decision-making on a shared "
"geospatial platform; and (3) monitoring of WASH-intervention effectiveness by tracking subsequent changes in environmental "
"risk indices. LGA-level infrastructure indicators provide a baseline for tracking progress toward WASH targets and identifying "
"investment priorities.")

H("5.4 Limitations",2)
P("This study acknowledges the following limitations. Environmental data constraints: optical indices (NDWI, NDVI) may be "
"obstructed by cloud cover during the peak rainy season; while Sentinel-1 SAR partially mitigates this, optical gaps remain, and "
"land-surface-temperature and humidity retrievals carry known uncertainties (Mobasheri, 2022). Data quality and completeness: "
"NCDC reporting is subject to delays, under-reporting, and variable diagnostic capacity, and national reporting is published at "
"state rather than LGA resolution, so a complete 774-LGA longitudinal panel cannot be constructed from official national reports "
"alone. Pilot scope: the Cross River demonstration is a small, single-state line-list used for technical validation and does not "
"support inferential epidemiological conclusions. Validation status: the risk-scoring algorithm uses fixed heuristic weights and "
"has not been prospectively validated against an independent, real-time outbreak dataset; the flood–cholera association is "
"exploratory and does not correct for temporal autocorrelation, spatial dependence, or multiple comparisons. Integration status: "
"live NCDC/SORMAS synchronization and multi-provider AI operation are described as design features and roadmap items, not as "
"evidenced production integrations. Technical capacity and sustainability: operation requires stable connectivity, GEE access, "
"and technical skills that may not be uniformly available in state health offices, and long-term sustainability depends on "
"institutional commitment, funding, and data-governance agreements.")

# ================= 6. FUTURE ROADMAP =================
H("6. Future Roadmap",1)
P("The evolution of the hub encompasses five priorities. First, socio-demographic vulnerability layer ingestion: overlaying "
"national health-infrastructure registries and Demographic and Health Survey (DHS) microdata—including access to safely managed "
"drinking water, improved sanitation, and household wealth quintiles—so that an environmental hazard triggers a high-risk alert "
"only where it intersects weak WASH infrastructure. Second, statistical calibration and validation: replacing fixed heuristic "
"weights with principal-component or outcome-calibrated weights, and evaluating the risk score on a temporal hold-out with "
"discrimination and calibration metrics, dependence-aware inference, and geographic leave-one-region-out sensitivity analysis. "
"Third, programmatic NCDC surveillance synchronization: establishing a governed API linkage (including SORMAS interoperability) "
"to replace manual extraction with automated case and death ingestion. Fourth, sovereign EO expansion: integrating NASRDA "
"historical archives and the planned four-satellite constellation (three optical, one SAR) to reduce dependence on foreign "
"data and improve cloud-penetrating flood mapping. Fifth, conversational AI for unstructured field reporting: extending the "
"assistant to parse field sitreps and community alerts into clean georeferenced records, with human verification, and expanding "
"multi-disease coverage to typhoid, dysentery, malaria, and dengue.")
figph("Figure 11")
caption("Figure 11: Advanced integration and automation roadmap — NCDC/SORMAS API, NASRDA satellite data hub, DHS microdata layer, "
"AI unstructured parser, and calibration/validation engine, unified through an integration gateway.")

# ================= 7. CONCLUSION =================
H("7. Conclusion",1)
P("This study presents the development and pilot demonstration of a scalable, Earth Observation-enabled environmental health "
"intelligence hub for cholera surveillance in Nigeria, together with a design for national scale-up. By harmonizing "
"satellite-derived environmental indices with officially reported epidemiological data, the platform provides an interpretable, "
"multi-agency decision-support environment. A Cross River line-list of 74 culture-referenced cases and 4 deaths across four LGAs "
"in 2021 was used to validate the ingestion, spatial-join, and risk-rendering workflow, while nationally reported NCDC figures—"
"led by the 111,062-case, 3,604-death 2021 epidemic—provide the national burden context. National reporting is published at the "
"state level; full LGA-level and ward-level resolution, statistical calibration, prospective forecasting, and live institutional "
"integration are identified as validation and governance priorities rather than completed results. Framed honestly against its "
"evidence base, the hub represents a credible technical foundation for transitioning Nigerian cholera surveillance from reactive "
"response toward proactive, environmentally contextualized, and sovereign public-health intelligence.")

# ================= 8. REFERENCES =================
H("8. References",1)
refs=[
"Adegoke, B. O., Odugbose, T., & Adeyemi, C. (2024). Data analytics for predicting disease outbreaks: A review of models and tools. International Journal of Life Science Research Updates, 2(2), 1-9.",
"Adesina, M. A., Adedeji, R. T., Oladipupo, I. R., Olufadewa, I. I., Oladele, R. I., Olufadewa, T. A., ... & Olansile, A. K. (2026). Cholera amid the climate crisis: a systematic review of flooding-driven health risks in West Africa. Discover Public Health, 23(1), 497.",
"Akanda, A. S., Jutla, A. S., Alam, M., De Magny, G. C., Siddique, A. K., Sack, R. B., ... & Islam, S. (2011). Hydroclimatic influences on seasonal and spatial cholera transmission cycles: implications for public health intervention in the Bengal Delta. Water Resources Research, 47(3).",
"Akingbola, A., Abiodun, A., Ojo, O., Jessica, O. U., Alao, U. H., Owolabi, A. O., & Chuku, J. (2025). Cholera outbreak in Nigeria: history, review of socioeconomic and meteorological drivers, diagnostic challenges, and artificial intelligence integration. Global Health, Epidemiology and Genomics, 2025(1), 8898076.",
"Albertini, C., Gioia, A., Iacobellis, V., & Manfreda, S. (2022). Detection of surface water and floods with multispectral satellites. Remote Sensing, 14(23), 6005.",
"Anderson, K., Ryan, B., Sonntag, W., Kavvada, A., & Friedl, L. (2017). Earth observation in service of the 2030 Agenda for Sustainable Development. Geo-spatial Information Science, 20(2), 77-96.",
"Armando, C. J., Rocklöv, J., Sidat, M., Tozan, Y., Mavume, A. F., Bunker, A., & Sewe, M. O. (2024). Spatial-temporal analysis of climate and socioeconomic conditions on cholera incidence in Mozambique from 2000 to 2018. BMJ Open, 14(8), e082503.",
"Atobatele, O. K., Ajayi, O. O., Hungbo, A. Q., & Adeyemi, C. (2019). Leveraging public health informatics to strengthen monitoring and evaluation of global health intervention. IRE Journals, 2(7), 174-193.",
"Avtar, R., Komolafe, A. A., Kouser, A., Singh, D., Yunus, A. P., Dou, J., ... & Kumar, P. (2020). Assessing sustainable development prospects through remote sensing: A review. Remote Sensing Applications: Society and Environment, 20, 100402.",
"Bhaga, T. D., Dube, T., Shekede, M. D., & Shoko, C. (2023). Investigating the effectiveness of Landsat-8 OLI and Sentinel-2 MSI satellite data in monitoring the effects of drought on surface water resources. Remote Sensing Applications: Society and Environment, 32, 101037.",
"Bhunia, G. S., & Shit, P. K. (2021). GeoComputation and Disease Ecology. In GeoComputation and Public Health: A Spatial Approach (pp. 151-220). Springer.",
"Bose, I., Hadida, G., Green, R., Murray, K. A., Part, C., & Kovats, S. (2026). Rainfall and water-related diseases, malnutrition and mortality in Low- and Middle-Income Countries: a systematic review. Heliyon, 12(1).",
"Brumfield, K. D., Usmani, M., Long, D. M., Lupari, H. A., Pope, R. K., Jutla, A. S., ... & Colwell, R. R. (2025). Climate change and Vibrio: Environmental determinants for predictive risk assessment. PNAS, 122(33), e2420423122.",
"Campbell, A. M., Racault, M. F., Goult, S., & Laurenson, A. (2020). Cholera risk: a machine learning approach applied to essential climate variables. International Journal of Environmental Research and Public Health, 17(24), 9378.",
"Chowdhury, F., Ross, A. G., Islam, M. T., McMillan, N. A. J., & Qadri, F. (2022). Diagnosis, Management, and Future Control of Cholera. Clinical Microbiology Reviews, 35(3), e0021121.",
"Chowdhury, A. H., & Rahman, M. S. (2025). Machine learning and spatio-temporal analysis of meteorological factors on waterborne diseases in Bangladesh. PLoS Neglected Tropical Diseases, 19(1), e0012800.",
"Christaki, E., Dimitriou, P., Pantavou, K., & Nikolopoulos, G. K. (2020). The impact of climate change on cholera: A review on the global status and future challenges. Atmosphere, 11(5), 449.",
"Cord, A. F., Brauman, K. A., Chaplin-Kramer, R., Huth, A., Ziv, G., & Seppelt, R. (2017). Priorities to advance monitoring of ecosystem services using earth observation. Trends in Ecology & Evolution, 32(6), 416-428.",
"da Silva Junior, U. J., da Penha Pacheco, A., Ruiz-Armenteros, A. M., et al. (2025). A methodological approach to flood dynamics based on satellite-derived spectral indices and altimetric forecast models. Science of the Total Environment, 1003, 180686.",
"Escobar, L. E., Ryan, S. J., Stewart-Ibarra, A. M., et al. (2015). A global map of suitability for coastal Vibrio cholerae under current and future climate conditions. Acta Tropica, 149, 202-211.",
"Ganesan, D., Gupta, S. S., & Legros, D. (2020). Cholera surveillance and estimation of burden of cholera. Vaccine, 38, 13-17.",
"Geng, Q., Yao, X., Sun, J., Jia, S., Xu, F., Zhang, L., ... & Li, G. (2026). Open-source GIS software and tools: a systematic review. Big Earth Data, 1-48.",
"Gopo, L., Bere, T., & Murisa, M. R. (2026). Climatic and socio-environmental drivers of cholera epidemics. PLOS Climate, 5(4), e0000840.",
"Gorelick, N., Hancher, M., Dixon, M., Ilyushchenko, S., Thau, D., & Moore, R. (2017). Google Earth Engine: Planetary-scale geospatial analysis for everyone. Remote Sensing of Environment, 202, 18-27.",
"Hashizume, M., Armstrong, B., Hajat, S., Wagatsuma, Y., Faruque, A. S. G., Hayashi, T., & Sack, D. A. (2008). The effect of rainfall on the incidence of cholera in Bangladesh. Epidemiology, 19(1), 103-110.",
"Huffman, G. J., Stocker, E. F., Bolvin, D. T., Nelkin, E. J., & Tan, J. (2020). GPM IMERG Final Precipitation L3 Half-Hourly 0.1° x 0.1°. NASA GES DISC.",
"Kamalrathne, T., Amaratunga, D., Haigh, R., & Kodituwakku, L. (2023). Need for effective detection and early warnings for epidemic and pandemic preparedness planning. International Journal of Disaster Risk Reduction, 92, 103724.",
"Kanagaraj, J. K., & Vijayan, T. B. (2024). Spatio-temporal correlation analysis of environmental and climatic determinants with various infectious diseases in Tamil Nadu and West Bengal. Environmental Engineering & Management Journal, 23(9).",
"Khachoo, Y. H., Cutugno, M., Robustelli, U., & Pugliano, G. (2026). Google Earth Engine since 2022: A structured bibliometric review of GeoAI-driven trends and applications. Sustainability, 18(12), 6241.",
"Kuna, A., & Gajewski, M. (2017). Cholera—the new strike of an old foe. International Maritime Health, 68(3), 163-167.",
"Li, X., Zhou, Z., Jia, H., Li, Z., Yang, Z., Cai, Z., ... & Shi, X. (2026). Mechanisms of accumulation–transport–discharge and source apportionment of combined sewer overflow pollution. Water, 18(5), 573.",
"Lloyd, C. T., Sturrock, H. J. W., Leasure, D. R., Jochem, W. C., Lázár, A. N., & Tatem, A. J. (2020). Using GIS and machine learning to classify residential status of urban buildings in low- and middle-income settings. Remote Sensing, 12(23), 3847.",
"Madubueze, M. H., Mbanefo, O. D., Anekwe, J. K., Nwadiogbu, N. M., & Egberi, A. E. (2025). Climate variability and waterborne disease burden in Nigeria. International Journal of Interdisciplinary Studies and Innovation, 4(1), 85-90.",
"McFeeters, S. K. (1996). The use of the Normalized Difference Water Index (NDWI) in the delineation of open water features. International Journal of Remote Sensing, 17(7), 1425-1432.",
"Meckawy, R., Stuckler, D., Mehta, A., Al-Ahdal, T., & Doebbeling, B. N. (2022). Effectiveness of early warning systems in the detection of infectious diseases outbreaks: a systematic review. BMC Public Health, 22(1), 2216.",
"Mikaberidze, A., et al. (2025). Opportunities and challenges in combining optical sensing and epidemiological modeling. Phytopathology, 115(10), 1260-1285.",
"Mobasheri, M. R. (2022). Remote Sensing and Computational Epidemiology. In COVID-19 Pandemic, Geospatial Information, and Community Resilience (pp. 55-68). CRC Press.",
"Morse, S. S., Mazet, J. A., Woolhouse, M., Parrish, C. R., Carroll, D., Karesh, W. B., ... & Daszak, P. (2012). Prediction and prevention of the next pandemic zoonosis. The Lancet, 380(9857), 1956-1965.",
"National Population Commission & ICF. (2024). Nigeria Demographic and Health Survey 2023-24: Key Indicators Report. Abuja & Rockville.",
"Nigeria Centre for Disease Control and Prevention (NCDC). (2022). Cholera Situation Report, Epi Week 52, 2021. Abuja: NCDC.",
"Nigeria Centre for Disease Control and Prevention (NCDC). (2024). Cholera Situation Report, Epi Week 22, 2024. Abuja: NCDC.",
"Nigeria Centre for Disease Control and Prevention (NCDC). (2025). An update of cholera outbreak in Nigeria. Federal Ministry of Health.",
"Olagunju, A. T., Denga, V. S., Madakadze, C., & Osuagwu, F. (2025). Climate change and health issues: an Afrocentric perspective. In Health and Climate Change (pp. 251-265). Academic Press.",
"Peprah, M. S., Moomen, A. W., & Sey, J. M. (2026). Leveraging Digital Earth for sustainable development in Africa: a systematic review and meta-analysis. Arabian Journal of Geosciences, 19(5), 99.",
"Pezanowski, S., Koua, E. L., Okeibunor, J. C., & Gueye, A. S. (2024). Predictors of disease outbreaks at continental scale in the African region. Digital Health, 10, 20552076241278939.",
"Reyburn, R., Kim, D. R., Emch, M., Khatib, A., Von Seidlein, L., & Ali, M. (2011). Climate variability and the outbreaks of cholera in Zanzibar, East Africa: a time series analysis. American Journal of Tropical Medicine and Hygiene, 84(6), 862.",
"Schets, F. M., Pol-Hofstad, I. E., van den Berg, H. H., & Schijven, J. F. (2025). Climate change-related temperature impact on human health risks of Vibrio species in bathing and surface water. Microorganisms, 13(8), 1893.",
"Singh, K., Kumar, S., Brumfield, K. D., Deliz, K., Colwell, R. R., Jutla, A., & Usmani, M. (2025). Remote Sensing to Predict Climate-Sensitive Pathogens: a Review. Authorea Preprints.",
"Singh, S., Sharma, P., Pal, N., Sarma, D. K., Tiwari, R., & Kumar, M. (2024). Holistic one health surveillance framework. ACS Infectious Diseases, 10(3), 808-826.",
"Song, Y., & Wu, P. (2021). Earth observation for sustainable infrastructure: A review. Remote Sensing, 13(8), 1528.",
"Tucker, C. J. (1979). Red and photographic infrared linear combinations for monitoring vegetation. Remote Sensing of Environment, 8(2), 127-150.",
"United Nations Children's Fund (UNICEF). (2025). Nigeria Humanitarian Flash Update (Cholera Outbreak), 1, 1-6.",
"Usmani, M., Brumfield, K. D., Magers, B. M., et al. (2023). Combating cholera by building predictive capabilities for pathogenic Vibrio cholerae in Yemen. Scientific Reports, 13(1), 2255.",
"Verma, R., & Kotwal, M. (2025). Geographic information systems: a review of its evolution, challenges, and future trends. International Journal for Multidisciplinary Research.",
"Warekuromor, T. (2026). Development and Pilot Validation of the Nigeria EO-enabled Environmental Health Intelligence Hub for Cholera Surveillance [Conference presentation]. GEO Health Community of Practice, NASRDA.",
"Wright, A. K. A., Ezugwu, C. I., Iregbu, J. K., et al. (2025). Climate change and emerging infectious diseases: a global review. Epidemiology and Health Data Insights, 1(3), ehdi009.",
"Yaghobi, S., Daneshi, A., Faramarzi, M., Azadi, H., Fathizad, H., & Islami, I. (2026). Drought impacts on vegetation cover in Western Iran. International Journal of Environmental Science and Technology, 23(1), 35.",
"Yu, H. (2025). Climate change unveils hidden microbial dangers. Environmental Science and Ecotechnology, 24, 100544.",
"Zhao, Q., Yu, L., Du, Z., Peng, D., Hao, P., Zhang, Y., & Gong, P. (2022). An overview of the applications of earth observation satellite data. Remote Sensing, 14(8), 1863.",
]
for r in sorted(refs):
    p=doc.add_paragraph(); run=p.add_run(r); run.font.size=Pt(9)
    p.paragraph_format.left_indent=Pt(18); p.paragraph_format.first_line_indent=Pt(-18)

doc.save('/root/CHOLERA_PAPER_V5_RealData.docx')
print("FULL PAPER saved. paragraphs:", len(doc.paragraphs), "| references:", len(refs))
